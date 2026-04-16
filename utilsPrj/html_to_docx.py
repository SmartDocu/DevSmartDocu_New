import re
import base64
import requests
import docx
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from utilsPrj.supabase_client import SUPABASE_SCHEMA

# ===== 유틸 =====
def hex_to_rgb(hex_color):
    if not hex_color: return (0,0,0)
    hex_color = hex_color.strip().lstrip('#')
    if len(hex_color) == 3: hex_color = ''.join([c*2 for c in hex_color])
    try: return tuple(int(hex_color[i:i+2],16) for i in (0,2,4))
    except: return (0,0,0)

def css_length_to_pt(val, base_pt=11.0, allow_relative=True):
    if not val: return None
    s = str(val).strip().lower()
    try:
        if s.endswith('pt'): return float(s[:-2])
        if s.endswith('px'): return float(s[:-2])*0.75
        if s.endswith('cm'): return float(s[:-2])*28.3464567
        if s.endswith('mm'): return float(s[:-2])*2.83464567
        if s.endswith('in'): return float(s[:-2])*72.0
        if allow_relative:
            if s.endswith('em'): return base_pt*float(s[:-2])
            if s.endswith('rem'): return base_pt*float(s[:-3])
            if s.endswith('%'): return base_pt*(float(s[:-1])/100)
        return float(s)
    except: return None

def css_length_to_inches(val):
    if not val: return None
    s = str(val).strip().lower()
    try:
        if s.endswith('in'): return float(s[:-2])
        if s.endswith('pt'): return float(s[:-2])/72.0
        if s.endswith('px'): return float(s[:-2])/96.0
        if s.endswith('cm'): return float(s[:-2])/2.54
        if s.endswith('mm'): return float(s[:-2])/25.4
        return None
    except: return None

def parse_css_style(style_str):
    styles = {}
    if not style_str: return styles
    for item in style_str.split(';'):
        if ':' in item:
            k,v = item.split(':',1)
            styles[k.strip().lower()]=v.strip()
    return styles

def resolve_font_size_pt(styles, default_pt=11.0):
    if not styles: return default_pt
    if '__font_size_pt' in styles: return styles['__font_size_pt']
    fs = styles.get('font-size')
    if fs:
        pt = css_length_to_pt(fs, base_pt=default_pt, allow_relative=True)
        if pt: return pt
    return default_pt

def merge_and_resolve_styles(parent_styles, child_styles):
    merged = {**(parent_styles or {}), **(child_styles or {})}
    if 'font-size' in child_styles:
        merged['__font_size_pt'] = resolve_font_size_pt(child_styles, 11.0)
    else:
        parent_pt = parent_styles.get('__font_size_pt', 11.0)
        merged['__font_size_pt'] = resolve_font_size_pt(merged, parent_pt)
    return merged

# ===== 텍스트 서식 =====
def set_run_font(run, font_name='굴림체'):
    if run:
        run.font.name = font_name
        r = run._element.rPr.rFonts
        r.set(qn('w:eastAsia'), font_name)

def apply_text_formatting(run, styles):
    if not styles: return
    size_pt = resolve_font_size_pt(styles, 11.0)
    run.font.size = Pt(size_pt)
    if 'color' in styles:
        c = styles['color'].strip()
        if c.startswith('#'): run.font.color.rgb = RGBColor(*hex_to_rgb(c))
    if 'font-weight' in styles:
        fw = styles['font-weight'].strip().lower()
        try:
            if 'bold' in fw or int(fw)>=600: 
                run.font.bold=True
        except:
            if 'bold' in fw: 
                run.font.bold=True
    if styles.get('font-style')=='italic': run.font.italic=True
    if 'text-decoration' in styles and 'underline' in styles['text-decoration']: run.font.underline=True
    set_run_font(run,'굴림체')

def apply_paragraph_formatting(paragraph, styles):
    if not styles: return
    align = styles.get('text-align','').strip().lower()
    if align=='center': paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif align=='right': paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    elif align=='justify': paragraph.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align=='left': paragraph.alignment=WD_ALIGN_PARAGRAPH.LEFT

# ===== 이미지 처리 =====
def add_image_to_doc(doc, img_element, styles):
    src = img_element.get('src')
    if not src:
        return
    try:
        width_in = css_length_to_inches(styles.get('width'))
        height_in = css_length_to_inches(styles.get('height'))
        image_data = None

        if src.startswith('data:image'):
            header, encoded = src.split(',', 1)
            image_data = BytesIO(base64.b64decode(encoded))
        elif src.startswith('http'):
            response = requests.get(src)
            if response.status_code == 200:
                image_data = BytesIO(response.content)
        else:
            with open(src, 'rb') as f:
                image_data = BytesIO(f.read())

        if image_data:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_data, width=Inches(width_in) if width_in else None,
                                          height=Inches(height_in) if height_in else None)
            apply_paragraph_formatting(paragraph, styles)
    except Exception as e:
        print(f"⚠️ 이미지 삽입 실패 ({src}): {e}")

# ===== 테이블 헤더 고정 =====
def set_table_header_repeat(table_row):
    """테이블 행을 헤더로 설정하여 매 페이지마다 반복되도록 함"""
    tr = table_row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = trPr.find(qn('w:tblHeader'))
    if tblHeader is None:
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)
    # print(f"✅ 테이블 헤더 반복 설정 완료 (행 인덱스: {table_row._tr.getparent().index(table_row._tr)})")

# ===== 테이블 처리 (개선됨) =====
def apply_cell_background(cell, styles):
    bg = styles.get('background-color','').strip()
    if bg.startswith('#'):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear')
        shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'),bg.lstrip('#'))
        tc_pr.append(shd)

def set_cell_margins(cell, top=None,left=None,bottom=None,right=None):
    def pt_to_twips(pt): return int(round(pt*20))
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    def set_one(side,val): 
        # print(f'Sied:{side} / val:{val}')
        if val is None: return
        el = tcMar.find(qn(f'w:{side}'))
        if el is None:
            el = OxmlElement(f'w:{side}')
            tcMar.append(el)
        el.set(qn('w:type'),'dxa')
        el.set(qn('w:w'),str(pt_to_twips(val)))
    set_one('top',top)
    set_one('left',left)
    set_one('bottom',bottom)
    set_one('right',right)

def parse_padding_to_pts(styles, base_pt=11.0):
    if not styles: return (None,None,None,None)
    if 'padding' in styles:
        parts = re.split(r'\s+',styles['padding'].strip())
        pts=[css_length_to_pt(p,base_pt=base_pt,allow_relative=False) for p in parts]
        if len(pts)==1: 
            t=b=l=r=pts[0]
        elif len(pts)==2: 
            t=b=pts[0]
            l=r=pts[1]
        elif len(pts)==3:
            t,lr,b=pts; l=r=lr
        elif len(pts)>=4:
            t,r,b,l=pts[:4]
        return (t,l,b,r)
    top=css_length_to_pt(styles.get('padding-top'),base_pt=base_pt,allow_relative=False)
    left=css_length_to_pt(styles.get('padding-left'),base_pt=base_pt,allow_relative=False)
    bottom=css_length_to_pt(styles.get('padding-bottom'),base_pt=base_pt,allow_relative=False)
    right=css_length_to_pt(styles.get('padding-right'),base_pt=base_pt,allow_relative=False)
    return (top,left,bottom,right)

def apply_column_width(cell, styles):
    w = styles.get('width')
    if not w: return
    try:
        if w.endswith('px'): cell.width=Cm(float(w[:-2])*0.0264583)
        elif w.endswith('cm'): cell.width=Cm(float(w[:-2]))
    except: pass

def apply_row_height(docx_row,height_in=None,exact=False):
    if not height_in: return
    docx_row.height=Inches(height_in) if height_in<1 else Pt(height_in)
    docx_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST

def add_table_to_doc(doc, table_element):
    # thead와 tbody 구분하여 처리
    thead = table_element.find('thead')
    tbody = table_element.find('tbody')
    
    if thead and tbody:
        # thead와 tbody가 명확히 구분된 경우
        header_rows = thead.find_all('tr')
        body_rows = tbody.find_all('tr')
        all_rows = header_rows + body_rows
        header_count = len(header_rows)
        # print(f"📋 thead/tbody 구조 감지: 헤더 {header_count}행, 본문 {len(body_rows)}행")
    else:
        # thead/tbody가 없는 경우, 모든 tr을 가져와서 th 태그로 헤더 판단
        all_rows = table_element.find_all('tr')
        header_rows = []
        header_count = 0
        
        # 연속된 th 태그를 포함한 행들을 헤더로 간주
        for row in all_rows:
            if row.find('th'):
                header_rows.append(row)
                header_count += 1
            else:
                break  # th가 없는 행이 나오면 헤더 끝
        
        # print(f"📋 th 태그 기반 헤더 감지: {header_count}행")
        
    if not all_rows: 
        print("⚠️ 테이블에 행이 없습니다.")
        return
    
    max_cols = max(len(row.find_all(['td','th'])) for row in all_rows)
    docx_table = doc.add_table(rows=len(all_rows), cols=max_cols)
    docx_table.style = 'Table Grid'
    docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    docx_table.autofit = False

    # 헤더 행들을 반복되도록 설정
    for i in range(header_count):
        set_table_header_repeat(docx_table.rows[i])
        # print(f"🔄 헤더 행 {i+1} 반복 설정 완료")

    # 컬럼 너비와 행 높이 계산
    col_widths = [0] * max_cols
    row_heights = [None] * len(all_rows)
    row_exact = [False] * len(all_rows)

    for r_idx, row in enumerate(all_rows):
        row_styles = parse_css_style(row.get('style', ''))
        h_css = row_styles.get('height')
        if h_css:
            h_in = css_length_to_inches(h_css)
            if h_in: 
                row_heights[r_idx] = h_in
                row_exact[r_idx] = True
        
        cells = row.find_all(['td', 'th'])
        for c_idx, cell in enumerate(cells[:max_cols]):
            cell_styles = parse_css_style(cell.get('style', ''))
            w_css = cell_styles.get('width') or cell.get('width')
            w_in = css_length_to_inches(w_css) if w_css else None
            if w_in and w_in > 0 and w_in > col_widths[c_idx]: 
                col_widths[c_idx] = w_in


    # 실제 셀 내용 처리
    for r_idx, row in enumerate(all_rows):
        cells = row.find_all(['td', 'th'])
        docx_row = docx_table.rows[r_idx]
        row_styles = parse_css_style(row.get('style', ''))
        
        if row_heights[r_idx]: 
            apply_row_height(docx_row, row_heights[r_idx], exact=row_exact[r_idx])

        for c_idx, cell in enumerate(cells[:max_cols]):
            docx_cell = docx_row.cells[c_idx]
            cell_styles = parse_css_style(cell.get('style', ''))
            merged_styles = merge_and_resolve_styles(row_styles, cell_styles)
            apply_cell_background(docx_cell, merged_styles)
            apply_column_width(docx_cell, cell_styles)
            
            t, l, b, r = parse_padding_to_pts(cell_styles, resolve_font_size_pt(merged_styles, 11.0))
            set_cell_margins(docx_cell, t, l, b, r)
            
            docx_cell.text = ""
            paragraph = docx_cell.paragraphs[0]
            process_text_content(paragraph, cell, merged_styles)
            apply_paragraph_formatting(paragraph, merged_styles)
            
            # th 태그이거나 헤더 행인 경우 굵게 처리
            if cell.name.lower() == 'th' or r_idx < header_count:
                for run in paragraph.runs: 
                    run.font.bold = True

# ===== 본문 콘텐츠 처리 =====
def process_text_content(paragraph, element, styles):
    """
    HTML element 안의 텍스트/태그를 Word paragraph에 추가하고
    CSS/태그 기반 서식을 적용한다.
    """
    for content in element.contents:
        if hasattr(content, 'name') and content.name:
            tag_name = content.name.lower()
            child_styles = parse_css_style(content.get('style', ''))
            merged_styles = merge_and_resolve_styles(styles, child_styles)

            # ---- 굵게 (strong, b) ----
            if tag_name in ['strong', 'b']:
                if content.contents:
                    for child in content.contents:
                        run = paragraph.add_run(child if isinstance(child, str) else child.get_text())
                        apply_text_formatting(run, merged_styles)
                        run.bold = True
                else:
                    run = paragraph.add_run(content.get_text())
                    apply_text_formatting(run, merged_styles)
                    run.bold = True

            # ---- 기울임 (em, i) ----
            elif tag_name in ['em', 'i']:
                if content.contents:
                    for child in content.contents:
                        run = paragraph.add_run(child if isinstance(child, str) else child.get_text())
                        apply_text_formatting(run, merged_styles)
                        run.italic = True
                else:
                    run = paragraph.add_run(content.get_text())
                    apply_text_formatting(run, merged_styles)
                    run.italic = True

            # ---- 밑줄 (u) ----
            elif tag_name == 'u':
                if content.contents:
                    for child in content.contents:
                        run = paragraph.add_run(child if isinstance(child, str) else child.get_text())
                        apply_text_formatting(run, merged_styles)
                        run.underline = True
                else:
                    run = paragraph.add_run(content.get_text())
                    apply_text_formatting(run, merged_styles)
                    run.underline = True

            # ---- 줄바꿈 (br) ----
            elif tag_name == 'br':
                paragraph.add_run().add_break()

            # ---- 일반 span, 링크 등 ----
            elif tag_name in ['span', 'a']:
                if content.contents:
                    process_text_content(paragraph, content, merged_styles)
                else:
                    run = paragraph.add_run(content.get_text())
                    apply_text_formatting(run, merged_styles)

            # ---- 기타 태그: 재귀 처리 ----
            else:
                process_text_content(paragraph, content, merged_styles)

        else:
            # 일반 텍스트 노드
            text = str(content)
            if text.strip():
                run = paragraph.add_run(text)
                apply_text_formatting(run, styles)

# # ===== 엘리먼트 처리 =====
# def process_element(doc, element, parent_styles=None):
#     if parent_styles is None: parent_styles={'__font_size_pt':11.0}
#     if isinstance(element,str):
#         text=element.strip()
#         if text:
#             p=doc.add_paragraph()
#             apply_paragraph_formatting(p,parent_styles)
#             run=p.add_run(text)
#             apply_text_formatting(run,parent_styles)
#         return
#     if not hasattr(element,'name') or element.name is None:
#         text=str(element).strip()
#         if text and text not in ['\n','\r\n','\t']:
#             p=doc.add_paragraph()
#             apply_paragraph_formatting(p,parent_styles)
#             run=p.add_run(text)
#             apply_text_formatting(run,parent_styles)
#         return
#     current_styles=parse_css_style(element.get('style',''))
#     combined_styles=merge_and_resolve_styles(parent_styles,current_styles)
#     tag_name=element.name.lower()

#     if tag_name=='img':
#         add_image_to_doc(doc, element, combined_styles)
#         return
#     if tag_name=='table':
#         add_table_to_doc(doc,element)
#         return
#     if tag_name in ['h1','h2','h3','h4','h5','h6']:
#         level=int(tag_name[1])
#         heading=doc.add_heading(level=level)
#         heading.text=""
#         process_text_content(heading,element,combined_styles)
#         apply_paragraph_formatting(heading,combined_styles)
#         return
#     if tag_name in ['p','div']:
#         p=doc.add_paragraph()
#         apply_paragraph_formatting(p,combined_styles)
#         process_text_content(p,element,combined_styles)
#         return
#     if tag_name in ['ul','ol']:
#         for li in element.find_all('li',recursive=False):
#             li_styles=parse_css_style(li.get('style',''))
#             merged=merge_and_resolve_styles(combined_styles,li_styles)
#             p=doc.add_paragraph(style='List Bullet' if tag_name=='ul' else 'List Number')
#             apply_paragraph_formatting(p,merged)
#             process_text_content(p,li,merged)
#         return
#     if tag_name=='br': doc.add_paragraph(); return
#     if tag_name in ['span','strong','b','em','i','u','a']: return
#     for child in element.children:
#         if hasattr(child,'name'): process_element(doc,child,combined_styles)
#         else:
#             text=str(child).strip()
#             if text:
#                 p=doc.add_paragraph()
#                 apply_paragraph_formatting(p,combined_styles)
#                 run=p.add_run(text)
#                 apply_text_formatting(run,combined_styles)

def process_element(doc, element, parent_styles=None):

    if parent_styles is None: parent_styles={'__font_size_pt':11.0}
    if isinstance(element,str):
        text=element.strip()
        if text:
            p=doc.add_paragraph()
            apply_paragraph_formatting(p,parent_styles)
            run=p.add_run(text)
            apply_text_formatting(run,parent_styles)
        return
    if not hasattr(element,'name') or element.name is None:
        text=str(element).strip()
        if text and text not in ['\n','\r\n','\t']:
            p=doc.add_paragraph()
            apply_paragraph_formatting(p,parent_styles)
            run=p.add_run(text)
            apply_text_formatting(run,parent_styles)
        return
    current_styles=parse_css_style(element.get('style',''))
    combined_styles=merge_and_resolve_styles(parent_styles,current_styles)
    
    tag_name = element.name.lower()

    if tag_name=='img':
        add_image_to_doc(doc, element, combined_styles)
        return
    if tag_name=='table':
        add_table_to_doc(doc,element)
        return
    if tag_name in ['h1','h2','h3','h4','h5','h6']:
        level=int(tag_name[1])
        heading=doc.add_heading(level=level)
        heading.text=""
        process_text_content(heading,element,combined_styles)
        apply_paragraph_formatting(heading,combined_styles)
        return
    # if tag_name in ['p','div']:
    #     p=doc.add_paragraph()
    #     apply_paragraph_formatting(p,combined_styles)
    #     process_text_content(p,element,combined_styles)
    #     return
    if tag_name == 'p':
        # 1️⃣ 이미지 먼저 처리
        imgs = element.find_all('img', recursive=True)
        for img in imgs:
            add_image_to_doc(doc, img, combined_styles)

        # 2️⃣ 테이블은 절대 text 처리하지 말고 그대로 넘김
        tables = element.find_all('table', recursive=True)
        for table in tables:
            process_element(doc, table, combined_styles)

        # 3️⃣ img, table 제거한 순수 텍스트만 문단으로 처리
        p=doc.add_paragraph()
        apply_paragraph_formatting(p,combined_styles)
        element_copy = element.__copy__()
        for tag in element_copy.find_all(['img', 'table']):
            tag.decompose()

        process_text_content(p, element_copy, combined_styles)
        return

    if tag_name in ['ul','ol']:
        for li in element.find_all('li',recursive=False):
            li_styles=parse_css_style(li.get('style',''))
            merged=merge_and_resolve_styles(combined_styles,li_styles)
            p=doc.add_paragraph(style='List Bullet' if tag_name=='ul' else 'List Number')
            apply_paragraph_formatting(p,merged)
            process_text_content(p,li,merged)
        return
    if tag_name=='br': doc.add_paragraph(); return
    if tag_name in ['span','strong','b','em','i','u','a']: return
    for child in element.children:
        if hasattr(child,'name'): process_element(doc,child,combined_styles)
        else:
            text=str(child).strip()
            if text:
                p=doc.add_paragraph()
                apply_paragraph_formatting(p,combined_styles)
                run=p.add_run(text)
                apply_text_formatting(run,combined_styles)






# ===== 최종 진입점 =====
def html_to_docx(supabase, genchapteruid, html_content):
    try:
        # print(f'HTML_Content: {html_content}')
        # 페이지 나누기 코드 변환 처리
        sep_pagebreak = '<div class="page-break" style="page-break-after:always;"><span style="display:none;">&nbsp;</span></div>'
        html_content = html_content.replace(sep_pagebreak, '---페이지 나누기---')

        soup=BeautifulSoup(html_content,'html.parser')
        
        doc=Document()
            
        set_headerfooter(supabase, genchapteruid, doc)
        for section in doc.sections:
            section.top_margin=Cm(2.54)
            section.bottom_margin=Cm(2.54)
            section.left_margin=Cm(2.54)
            section.right_margin=Cm(2.54)

        body=soup.find('body')
        elements_to_process=body.find_all(recursive=False) if body else [e for e in soup.children if hasattr(e,'name')]
        for element in elements_to_process:
            process_element(doc,element,parent_styles={'__font_size_pt':11.0})

        # 혹시라도 HTML에서 구조적으로 변환된 문단이 하나도 없을 경우, 최소한의 텍스트라도 문서에 넣자
        if len(doc.paragraphs)==0:
            t=soup.get_text(strip=True)
            if t:
                p=doc.add_paragraph(t)

        # 페이지 나누기 코드를 워드 페이지 나누기로 변환 --> Text To Word
        target_text = '---페이지 나누기---'
        for i, paragraph in enumerate(doc.paragraphs):
            if target_text in paragraph.text:
                paragraph.text = paragraph.text.replace(target_text, "")
                
                # 페이지 나누기
                run = paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
                break

        # print(f"✅ 변환 완료")
        for para in doc.paragraphs:
            # 줄 간격 1.0 설정
            para.paragraph_format.line_spacing = 1.0  

        # print('HTML_To_Docx 변환 작업 완료')
        return doc
    except Exception as e:
        print(f"❌ 변환 중 오류 발생: {e}")
        return doc

# 머리글 바닥글 설정
def set_headerfooter(supabase, genchapteruid, result_doc):
    # chapter 조회
    chapteruid = supabase.schema(SUPABASE_SCHEMA).table('genchapters').select('chapteruid').eq('genchapteruid', genchapteruid).execute().data[0]['chapteruid']
    # 챕터 템플릿 구하기
    chapter = supabase.schema(SUPABASE_SCHEMA).table('chapters').select('docid', 'chapternm', 'chaptertemplateurl').eq('chapteruid', chapteruid).execute().data
    docid = chapter[0]['docid']
    chaptertemplate_url = chapter[0]['chaptertemplateurl']
    # 문서 템플릿 구하기
    basetemplateurl = supabase.schema(SUPABASE_SCHEMA).table('docs').select('basetemplateurl').eq('docid', docid).execute().data[0]['basetemplateurl']

    templateyn = False
    # 챕터 템플릿이 있으면 해당 사항 / 그 외 기본 템플릿
    if chaptertemplate_url:
        response = requests.get(chaptertemplate_url)
        templateyn = True
    elif basetemplateurl:
        response = requests.get(basetemplateurl)
        templateyn = True
    
    if templateyn:
        template_doc = docx.Document(BytesIO(response.content))
        # 섹션별로 접근
        section = template_doc.sections[0]
        header = section.header.paragraphs[0].text
        footer = section.footer.paragraphs[0].text

        # 결과 파일에 머리글 바닥글 삽입
        result_section = result_doc.sections[0]
        result_section_header = result_section.header
        result_section_footer = result_section.footer
        result_section_header.paragraphs[0].text = header
        result_section_footer.paragraphs[0].text = footer
    else:
        pass

    return result_doc
    
# 병합용 -> 순서대로 
def html_to_docx_merge(supabase, doc, genchapteruid, html_content, index, previous_yn, current_yn):
    """
    Supabase 챕터/템플릿 기반으로 머리글·바닥글 적용, HTML 본문 병합
    index: 챕터 순서 (1부터 시작)
    """
    try:
        # 페이지 나누기 코드 변환 처리 --> HTML To Text
        sep_pagebreak = '<div class="page-break" style="page-break-after:always;"><span style="display:none;">&nbsp;</span></div>'
        html_content = html_content.replace(sep_pagebreak, '---페이지 나누기---')
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 1️⃣ 새 섹션 생성 + 머리글/바닥글 설정
        doc, previous_yn, current_yn = set_headerfooter_merge(supabase, genchapteruid, doc, index, previous_yn, current_yn)
        current_section = doc.sections[-1]
        
        # 2️⃣ 새 섹션을 독립적으로 만들기 (이전 섹션과 연결 끊기)
        current_section.header.is_linked_to_previous = False
        current_section.footer.is_linked_to_previous = False
        
        # 3️⃣ 여백/마진 적용
        current_section.top_margin = Cm(2.54)
        current_section.bottom_margin = Cm(2.54)
        current_section.left_margin = Cm(2.54)
        current_section.right_margin = Cm(2.54)
        
        # 4️⃣ HTML 본문 추가
        body = soup.find('body')
        elements_to_process = body.find_all(recursive=False) if body else [e for e in soup.children if hasattr(e,'name')]
        for element in elements_to_process:
            process_element(doc, element, parent_styles={'__font_size_pt': 11.0})

        # 5️⃣ HTML에서 문단이 없는 경우 최소 텍스트 추가
        if len(doc.paragraphs) == 0:
            t = soup.get_text(strip=True)
            if t:
                p = doc.add_paragraph(t)
        
        # 페이지 나누기 코드를 워드 페이지 나누기로 변환 --> Text To Word
        target_text = '---페이지 나누기---'
        for i, paragraph in enumerate(doc.paragraphs):
            if target_text in paragraph.text:
                paragraph.text = paragraph.text.replace(target_text, "")
                
                # 페이지 나누기
                run = paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
                break
        
        # 6️⃣ 줄 간격 1.0 설정
        for para in doc.paragraphs:
            para.paragraph_format.line_spacing = 1.0

        # 2025-12-15 Min 추가
        doc.add_page_break()
        
        # print(f"✅ 챕터 {index} 병합 완료")
        return previous_yn, current_yn

    except Exception as e:
        print(f"❌ 챕터 {index} 병합 중 오류 발생: {e}")
        return False, False

# 머리글 바닥글 설정 (병합용) - 수정된 버전
def set_headerfooter_merge(supabase, genchapteruid, result_doc, index, previous_yn, current_yn):
    """
    챕터별로 템플릿 머리글/바닥글 적용 + 새 섹션 생성
    """
    which = index - 1

    # chapter 조회
    chapteruid = supabase.schema(SUPABASE_SCHEMA).table('genchapters').select('chapteruid') \
        .eq('genchapteruid', genchapteruid).execute().data[0]['chapteruid']

    chapter = supabase.schema(SUPABASE_SCHEMA).table('chapters').select('docid', 'chapternm', 'chaptertemplateurl') \
        .eq('chapteruid', chapteruid).execute().data
    docid = chapter[0]['docid']
    chapternm = chapter[0]['chapternm']
    chaptertemplate_url = chapter[0]['chaptertemplateurl']

    basetemplateurl = supabase.schema(SUPABASE_SCHEMA).table('docs').select('basetemplateurl') \
        .eq('docid', docid).execute().data[0]['basetemplateurl']

    if chaptertemplate_url:
        current_yn = True
    else:
        current_yn = False

    # print(f"Index: {index} / CurrentYN: {current_yn} / Previous_yn: {previous_yn}")
    # 첫 섹션이면 기존 섹션 사용, 아니면 새 섹션 추가
    if which == 0:
        result_section = result_doc.sections[0]
    elif current_yn:
        # print('현재 챕터가 ChapterTempl')
        result_section = result_doc.add_section(WD_SECTION.NEW_PAGE)
        # 🔹 이전 섹션과 머리글/바닥글 연결 끊기
        result_section.header.is_linked_to_previous = False
        result_section.footer.is_linked_to_previous = False
    elif not(current_yn) and previous_yn:
        # print(f"현재 챕터는 BaseTempl 이지만 이전이 ChapterTempl")
        result_section = result_doc.add_section(WD_SECTION.NEW_PAGE)
        # 🔹 이전 섹션과 머리글/바닥글 연결 끊기
        result_section.header.is_linked_to_previous = False
        result_section.footer.is_linked_to_previous = False
    else:
        # print('현재 챕터 BaseTempl')
        result_section = result_doc.sections[0]
        # 🔹 이전 섹션과 머리글/바닥글 연결 끊기
        result_section.header.is_linked_to_previous = False
        result_section.footer.is_linked_to_previous = False

    # 템플릿 가져오기
    template_url = chaptertemplate_url or basetemplateurl

    if template_url:
        response = requests.get(template_url)
        template_doc = docx.Document(BytesIO(response.content))
        template_section = template_doc.sections[0]
        # 머리글/바닥글 복사
        if template_section.header.paragraphs:
            text = template_section.header.paragraphs[0].text
            if result_section.header.paragraphs:
                result_section.header.paragraphs[0].text = text
            else:
                result_section.header.add_paragraph().text = text

        if template_section.footer.paragraphs:
            text = template_section.footer.paragraphs[0].text
            if result_section.footer.paragraphs:
                result_section.footer.paragraphs[0].text = text
            else:
                result_section.footer.add_paragraph().text = text

    previous_yn = current_yn

    return result_doc, previous_yn, current_yn


def copy_header_footer_from_template(template_section, target_section):
    """템플릿 섹션에서 타겟 섹션으로 머리글/바닥글을 안전하게 복사"""
    
    try:
        # 머리글 복사
        template_header = template_section.header
        target_header = target_section.header
        
        # 템플릿에서 머리글 텍스트 가져오기
        header_text = ""
        if template_header.paragraphs:
            header_text = template_header.paragraphs[0].text
        
        # 타겟 머리글에 설정 (기존 문단 재사용)
        if target_header.paragraphs:
            target_header.paragraphs[0].text = header_text
        else:
            para = target_header.add_paragraph()
            para.text = header_text
        
        # 바닥글 복사
        template_footer = template_section.footer  
        target_footer = target_section.footer
        
        # 템플릿에서 바닥글 텍스트 가져오기
        footer_text = ""
        if template_footer.paragraphs:
            footer_text = template_footer.paragraphs[0].text
            
        # 타겟 바닥글에 설정 (기존 문단 재사용)
        if target_footer.paragraphs:
            target_footer.paragraphs[0].text = footer_text
        else:
            para = target_footer.add_paragraph()
            para.text = footer_text
        
        # print(f"머리글: '{header_text}', 바닥글: '{footer_text}' 설정 완료")
        
    except Exception as e:
        print(f"⚠️ 머리글/바닥글 복사 중 오류: {e}")


# 대안: 더 간단하고 안전한 방법
def set_headerfooter_merge_safe(supabase, genchapteruid, result_doc, index):
    """안전한 버전의 머리글/바닥글 설정"""
    which = index - 1
    
    # chapter 조회
    chapteruid = supabase.schema(SUPABASE_SCHEMA).table('genchapters').select('chapteruid').eq('genchapteruid', genchapteruid).execute().data[0]['chapteruid']
    
    # 챕터 템플릿 구하기
    chapter = supabase.schema(SUPABASE_SCHEMA).table('chapters').select('docid', 'chapternm', 'chaptertemplateurl').eq('chapteruid', chapteruid).execute().data
    docid = chapter[0]['docid']
    chapternm = chapter[0]['chapternm']
    chaptertemplate_url = chapter[0]['chaptertemplateurl']
    
    # 문서 템플릿 구하기
    basetemplateurl = supabase.schema(SUPABASE_SCHEMA).table('docs').select('basetemplateurl').eq('docid', docid).execute().data[0]['basetemplateurl']

    # 챕터 템플릿이 있으면 해당 사항 / 그 외 기본 템플릿
    if chaptertemplate_url:
        response = requests.get(chaptertemplate_url)
    else:
        response = requests.get(basetemplateurl)
    
    template_doc = docx.Document(BytesIO(response.content))
    
    # 템플릿에서 머리글/바닥글 텍스트 추출
    header_text = ""
    footer_text = ""
    if template_doc.sections[0].header.paragraphs:
        header_text = template_doc.sections[0].header.paragraphs[0].text
    if template_doc.sections[0].footer.paragraphs:
        footer_text = template_doc.sections[0].footer.paragraphs[0].text
    
    # print(f'템플릿에서 추출 - Header: "{header_text}", Footer: "{footer_text}"')
    
    if which == 0:
        # 첫 번째 섹션
        result_section = result_doc.sections[0]
        # print(f'첫 번째 섹션 설정')
    else:
        # 새로운 섹션 추가
        # print(f'새로운 섹션 생성: {index}번째')
        result_section = result_doc.add_section(WD_SECTION.NEW_PAGE)
    
    # 머리글/바닥글 설정 (안전한 방법)
    try:
        if result_section.header.paragraphs:
            result_section.header.paragraphs[0].text = header_text
        else:
            result_section.header.add_paragraph().text = header_text
            
        if result_section.footer.paragraphs:
            result_section.footer.paragraphs[0].text = footer_text  
        else:
            result_section.footer.add_paragraph().text = footer_text
            
        # print(f"섹션 {index} 머리글/바닥글 설정 완료")
        
    except Exception as e:
        print(f"⚠️ 섹션 {index} 머리글/바닥글 설정 실패: {e}")
    
    return result_doc


def disconnect_section_headers_footers(section):
    """새 섹션의 머리글/바닥글을 이전 섹션과 안전하게 분리"""
    try:
        # 가장 안전한 방법: 기존 내용만 지우고 새로 작성
        header = section.header
        footer = section.footer
        
        # 기존 문단들의 텍스트만 초기화 (구조는 유지)
        for paragraph in header.paragraphs:
            paragraph.clear()
        
        for paragraph in footer.paragraphs:
            paragraph.clear()
        
        # 문단이 없다면 새로 생성
        if not header.paragraphs:
            header.add_paragraph()
        if not footer.paragraphs:
            footer.add_paragraph()
            
        # print(f"섹션 머리글/바닥글 안전하게 초기화 완료")
        return section
        
    except Exception as e:
        print(f"⚠️ 섹션 분리 중 오류 (안전모드로 진행): {e}")
        return section


# 더 간단한 대안 방법
def create_independent_section_simple(doc):
    """간단한 방법: 새 문서를 만들고 내용을 복사하는 방식"""
    # 새로운 임시 문서 생성
    temp_doc = Document()
    
    # 현재 문서의 마지막 섹션 가져오기
    last_section = doc.sections[-1]
    
    # 새 섹션 추가 (완전히 독립적)
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # 섹션 속성 초기화
    new_section._sectPr.clear()
    
    # 기본 섹션 속성 다시 설정
    from docx.oxml.shared import OxmlElement, qn
    
    # 페이지 크기 설정
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')  # A4 width
    pgSz.set(qn('w:h'), '16838')  # A4 height
    new_section._sectPr.append(pgSz)
    
    # 페이지 여백 설정  
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '1440')
    pgMar.set(qn('w:right'), '1440') 
    pgMar.set(qn('w:bottom'), '1440')
    pgMar.set(qn('w:left'), '1440')
    new_section._sectPr.append(pgMar)
    
    return new_section