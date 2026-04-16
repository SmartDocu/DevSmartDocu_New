import base64
import requests
from io import BytesIO
import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_text_from_docx(docs, url):
    """Word 문서에서 순수 텍스트만 추출"""
    if url:
        response = requests.get(docs)
        doc = docx.Document(BytesIO(response.content))
    else:
        doc = docs

    try:
        if doc:
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return "⚠️ 문서를 불러올 수 없습니다. 파일이 존재하지 않거나, 다운로드 오류가 발생했습니다."
    except requests.exceptions.RequestException as e:
        return f"⚠️ 네트워크 오류가 발생했습니다: {e}"
    except Exception as e:
        return f"⚠️ 문서 처리 중 오류가 발생했습니다: {e}"


def convert_docx_to_html_2(docs, url, formyn, ckeditor_mode=False, printyn=False):
    """Word 문서를 HTML로 변환하는 메인 함수 (이미지 처리 포함)"""
    # formyn -> HTML 형식이 필요한가??
    # printyn 출력 용인지 보여주기 용인지??
    if url:
        response = requests.get(docs)
        doc = docx.Document(BytesIO(response.content))
    else:
        doc = docs
    
    # 섹션별로 접근
    for i, section in enumerate(doc.sections):
        header = section.header
        footer = section.footer
        
        # print(f"\n=== Section {i+1} ===")
        
        # 머릿글
        # for para in header.paragraphs:
        #     print("Header:", para.text)
        
        # 바닥글
        # for para in footer.paragraphs:
        #     print("Footer:", para.text)

    # HTML 기본 구조 설정
    html_parts = _get_html_header(formyn, ckeditor_mode)
    
    # 이미지 추출
    image_dict = _extract_images_from_docx(doc)
    
    # 문서 내용 처리 (단락과 표를 순서대로)
    current_list = None
    list_items = []
    
    for element in doc.element.body:
        if isinstance(element, CT_P):  # 단락 처리
            para = Paragraph(element, doc)
            html_content = _process_paragraph(para, image_dict, current_list, list_items, printyn)
            
            if html_content['close_list']:
                html_parts.append(_process_list(current_list, list_items))
                current_list = None
                list_items = []
            
            if html_content['new_list']:
                current_list = html_content['list_type']
                list_items = [html_content['content']]
            elif html_content['continue_list']:
                list_items.append(html_content['content'])
            else:
                html_parts.append(html_content['content'])
                
        elif isinstance(element, CT_Tbl):  # 표 처리
            if current_list:  # 표 시작 전에 열린 리스트 닫기
                html_parts.append(_process_list(current_list, list_items))
                current_list = None
                list_items = []
            
            table = Table(element, doc)
            table_html = _process_table(table, image_dict, ckeditor_mode, printyn)
            html_parts.append(table_html)
    
    # 열린 리스트가 있으면 닫기
    if current_list:
        html_parts.append(_process_list(current_list, list_items))
    
    # HTML 마무리
    if formyn:
        html_parts.extend(['</body>', '</html>'])
    
    return '\n'.join(html_parts)


def _get_html_header(formyn, ckeditor_mode=False):
    """HTML 헤더 생성"""
    if formyn:
        css_styles = [
            '.left { text-align: left; }',
            '.center { text-align: center; }',
            '.right { text-align: right; }',
            '.justify { text-align: justify; }',
            'table { border-collapse: collapse; width: 100%; margin: 10px 0; page-break-inside: auto; }',
            'thead { display: table-header-group !important; page-break-inside: avoid !important; }',
            'tbody { display: table-row-group; }',
            'table, th, td { border: 1px solid black; padding: 4px; }',
            'th { page-break-inside: avoid; mso-header-data: yes; }',  # 워드 전용 속성
            'tr:first-child { page-break-inside: avoid; page-break-after: avoid; }',  # 첫 번째 행 보호
            '.image-container { margin: 10px 0; text-align: center; }',
            '.image-container img { max-width: 100%; height: auto; }',
            # 워드 변환을 위한 추가 CSS
            '@media print { thead { display: table-header-group !important; } }',
            '@page { margin: 1in; }'
        ]
        
        if ckeditor_mode:
            # CKEditor 5 호환 스타일 추가
            css_styles.extend([
                '.text-center { text-align: center !important; }',
                '.text-right { text-align: right !important; }',
                '.text-justify { text-align: justify !important; }',
                '.cell-bg-color { background-color: var(--cell-bg-color) !important; }',
                # CKEditor에서 스타일이 제거되는 것을 방지하기 위한 강제 적용
                'th[style*="background-color"], td[style*="background-color"] { background-color: inherit !important; }'
            ])
        
        return [
            '<!DOCTYPE html>', '<html>', '<head>', 
            '<meta charset="UTF-8">', 
            '<style>',
            *css_styles,
            '</style>',
            '</head>', '<body>'
        ]
    else:
        return ['']


def _extract_images_from_docx(doc):
    """워드 문서에서 모든 이미지를 추출하여 딕셔너리로 반환"""
    image_dict = {}
    
    try:
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    # 이미지 데이터 읽기
                    image_data = rel.target_part.blob
                    
                    # MIME 타입 결정
                    content_type = rel.target_part.content_type
                    if not content_type:
                        filename = rel.target_part.partname
                        if filename.endswith('.png'):
                            content_type = 'image/png'
                        elif filename.endswith(('.jpg', '.jpeg')):
                            content_type = 'image/jpeg'
                        elif filename.endswith('.gif'):
                            content_type = 'image/gif'
                        else:
                            content_type = 'image/png'
                    
                    # base64 인코딩
                    base64_image = base64.b64encode(image_data).decode('utf-8')
                    image_dict[rel_id] = {
                        'data': f"data:{content_type};base64,{base64_image}",
                        'content_type': content_type
                    }
                except Exception as e:
                    print(f"이미지 처리 중 오류 발생: {e}")
                    continue
    except Exception as e:
        print(f"이미지 추출 중 오류 발생: {e}")
    
    return image_dict

def is_page_break_paragraph(paragraph):
    """단락 안에 페이지 브레이크가 포함되어 있는지 확인"""
    for br in paragraph._element.xpath('.//w:br'):
        if br.get(qn('w:type')) == 'page':
            return True
    return False

def _process_paragraph(paragraph, image_dict, current_list, list_items, printyn):
    """단락 처리 및 리스트 상태 관리"""
    result = {
        'content': '',
        'new_list': False,
        'continue_list': False,
        'close_list': False,
        'list_type': None
    }
    
    
    # 페이지 나누기를 코드로 --> Word To HTML
    if is_page_break_paragraph(paragraph):
        # result['content'] += '<div style="page-break-after: always; break-after: page; mso-break-type: page-break;"></div>'
        result['content'] += '<div class="page-break ck-widget" contenteditable="false"><span class="page-break__label">페이지 나누기</span></div>'
        return result

    # 빈 단락 처리
    if not paragraph.text.strip() and not _has_images_in_paragraph(paragraph):
        if current_list:
            result['close_list'] = True
        result['content'] = '<p>&nbsp;</p>'
        return result
    
    # 목록 스타일 확인
    style_name = paragraph.style.name.lower()
    if "list" in style_name:
        list_type = "ul" if "bullet" in style_name else "ol"
        content = _process_paragraph_text_with_images(paragraph, image_dict, printyn)
        
        if not current_list:
            result['new_list'] = True
            result['list_type'] = list_type
        else:
            result['continue_list'] = True
        result['content'] = content
        return result
    elif current_list:
        result['close_list'] = True
    
    # 제목 스타일 확인
    if style_name.startswith("heading"):
        try:
            level = int(style_name.split()[-1])
            if 1 <= level <= 6:
                alignment = _get_alignment_class(paragraph.alignment)
                content = _process_paragraph_text_with_images(paragraph, image_dict, printyn)
                result['content'] = f'<h{level}{alignment}>{content}</h{level}>'
                return result
        except:
            pass

    # 일반 단락 처리
    alignment = _get_alignment_class(paragraph.alignment)
    indent_style = _get_indent_styles(paragraph)
    style_attr = f' style="{";".join(indent_style)}"' if indent_style else ''
    content = _process_paragraph_text_with_images(paragraph, image_dict, printyn)
    result['content'] = f'<p{alignment}{style_attr}>{content}</p>'
    
    return result


def _has_images_in_paragraph(paragraph):
    """단락에 이미지가 포함되어 있는지 확인"""
    try:
        for run in paragraph.runs:
            if run.element.xpath('.//a:blip'):
                return True
    except:
        pass
    return False


def _get_alignment_class(alignment):
    """정렬 방식을 CSS 스타일로 변환"""
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return ' style="text-align:center;"'
    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return ' style="text-align:right;"'
    elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return ' style="text-align:justify;"'
    return ''


def _get_indent_styles(paragraph):
    """들여쓰기 스타일 추출"""
    indent_style = []
    
    try:
        if paragraph.paragraph_format.first_line_indent:
            px = paragraph.paragraph_format.first_line_indent.inches * 96
            indent_style.append(f'text-indent:{px:.0f}px')
        
        if paragraph.paragraph_format.left_indent:
            px = paragraph.paragraph_format.left_indent.inches * 96
            indent_style.append(f'margin-left:{px:.0f}px')
        
        if paragraph.paragraph_format.right_indent:
            px = paragraph.paragraph_format.right_indent.inches * 96
            indent_style.append(f'margin-right:{px:.0f}px')
    except:
        pass
    
    return indent_style


def _process_paragraph_text_with_images(paragraph, image_dict, printyn):
    """단락의 텍스트와 이미지를 HTML로 변환"""
    result = []
    current_width_px = 0

    # 탭 정지 위치 추출
    tab_stops = []
    try:
        if paragraph.paragraph_format.tab_stops:
            for stop in paragraph.paragraph_format.tab_stops:
                if stop.position:
                    tab_stops.append(_emu_to_px(stop.position))
        tab_stops.sort()
    except:
        pass

    for run in paragraph.runs:
        # 이미지 처리
        try:
            blip_elements = run.element.xpath('.//a:blip')
            for blip in blip_elements:
                embed_attr = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_attr and embed_attr in image_dict:
                    image_data = image_dict[embed_attr]['data']
                    
                    # 이미지 크기 정보 가져오기
                    size_attr = _get_image_size_attr(run)
                    result.append(f'<div class="image-container"><img src="{image_data}"{size_attr} alt="Document Image" /></div>')
        except:
            pass
        
        # 텍스트 처리
        processed_text = _process_run_text(run, tab_stops, current_width_px)
        if processed_text:
            styled_text = _apply_text_styles(run, processed_text, printyn)
            result.append(styled_text)
    
    return ''.join(result)


def _get_image_size_attr(run):
    """이미지 크기 속성 추출"""
    try:
        drawing = run.element.xpath('.//wp:drawing')[0]
        extent = drawing.xpath('.//wp:extent')[0]
        width_emu = int(extent.get('cx', '0'))
        height_emu = int(extent.get('cy', '0'))
        
        width_px = _emu_to_px(width_emu)
        height_px = _emu_to_px(height_emu)
        
        return f' width="{width_px}" height="{height_px}"'
    except:
        return ''


def _process_run_text(run, tab_stops, current_width_px):
    """Run의 텍스트 처리 (탭, 줄바꿈 포함)"""
    processed_text = ''
    
    try:
        run_xml = run._element
        for node in run_xml.iter():
            if node.tag == qn('w:tab'):
                # 탭 처리
                next_tab = None
                for stop in tab_stops:
                    if stop > current_width_px:
                        next_tab = stop
                        break
                tab_width = next_tab - current_width_px if next_tab else 50
                processed_text += f'<span style="display:inline-block; width:{tab_width}px;"></span>'
                current_width_px += tab_width
                
            elif node.tag == qn('w:t') and node.text:
                text = node.text
                # HTML 특수문자 이스케이프
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                # 줄바꿈과 공백 처리
                text = text.replace('\n', '<br>')
                text = text.replace('  ', '&nbsp;&nbsp;')
                processed_text += text
                # 글자 길이 누적 (대략 계산)
                current_width_px += len(text) * 10
                
            elif node.tag == qn('w:br'):
                processed_text += '<br>'
    except:
        # 기본 텍스트 처리 (오류 발생시)
        if run.text:
            text = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            text = text.replace('\n', '<br>').replace('  ', '&nbsp;&nbsp;')
            processed_text = text
    
    return processed_text


def _apply_text_styles(run, text, printyn):
    """텍스트에 스타일 적용 (굵기, 기울임, 밑줄, 취소선, 위/아래첨자, 색상 포함)"""
    if not text:
        return text

    style_attrs = []
    try:
        # --- 글자 크기 ---
        if run.font.size:
            pt_size = run.font.size.pt * 1.33  # pt → px (96dpi 기준)
            style_attrs.append(f'font-size:{int(pt_size)}px')

        # --- 글꼴 이름 ---
        if run.font.name:
            if run.font.name == '굴림체' and not(printyn):
                fontname = '굴림체, GulimChe, sans-serif'
            else:
                fontname = run.font.name
            style_attrs.append(f"font-family:'{fontname}'")

        # --- 글자 색상 ---
        if run.font.color and run.font.color.rgb:
            hex_color = f"#{run.font.color.rgb}"  # RGB → HEX (#RRGGBB)
            style_attrs.append(f"color:{hex_color}")

    except Exception as e:
        print(f"스타일 속성 처리 중 오류: {e}")

    style_attr = f' style="{";".join(style_attrs)}"' if style_attrs else ''

    # --- 굵기 상속 여부 확인 ---
    def is_bold_from_style(paragraph_style):
        try:
            style = paragraph_style
            while style:
                if hasattr(style, "font") and style.font.bold:
                    return True
                style = getattr(style, "based_on", None)
        except:
            pass
        return False

    is_bold = False
    try:
        if run.bold is True:
            is_bold = True
        elif run.bold is None:
            if is_bold_from_style(run._parent.style):
                is_bold = True
            if run.style.font.bold:
                is_bold = True
    except:
        pass

    # --- 기울임 여부 ---
    is_italic = False
    try:
        if run.italic is True:
            is_italic = True
    except:
        pass

    # --- 텍스트 스타일 적용 ---
    if is_bold and is_italic:
        text = f'<span{style_attr}><strong><em>{text}</em></strong></span>'
    elif is_bold:
        text = f'<span{style_attr}><strong>{text}</strong></span>'
    elif is_italic:
        text = f'<em{style_attr}>{text}</em>'
    elif style_attrs:
        text = f'<span{style_attr}>{text}</span>'

    # --- 추가 효과 ---
    try:
        if run.underline:
            text = f'<u>{text}</u>'
        if hasattr(run.font, 'strike') and run.font.strike:
            text = f'<s>{text}</s>'
        if hasattr(run.font, 'superscript') and run.font.superscript:
            text = f'<sup>{text}</sup>'
        elif hasattr(run.font, 'subscript') and run.font.subscript:
            text = f'<sub>{text}</sub>'
    except:
        pass

    return text



def _process_table(table, image_dict, ckeditor_mode=False, printyn=False):
    """표를 HTML로 변환 (색상 포함, 헤더 반복 설정)"""
    # 워드 변환을 위한 추가 속성들
    table_attrs = [
        'border-collapse: collapse',
        'width: 100%',
        'page-break-inside: auto'
    ]
    
    # 워드 전용 MSO 속성 추가
    mso_attrs = [
        'mso-table-layout-alt: fixed',
        'mso-table-wrap: around',
        'mso-table-lspace: 0pt',
        'mso-table-rspace: 0pt'
    ]
    
    table_style = '; '.join(table_attrs + mso_attrs)
    table_html = [f'<table style="{table_style}">']
    
    try:
        if table.rows:
            # 헤더 섹션 시작 (워드에서 인식하도록)
            table_html.append('<thead style="display: table-header-group; page-break-inside: avoid;">')
            
            # 헤더 행 처리 - 워드 전용 속성 추가
            header_row_style = [
                'page-break-inside: avoid',
                'page-break-after: avoid',
                'mso-row-height-rule: exactly'
            ]
            table_html.append(f'<tr style="{"; ".join(header_row_style)}">')
            
            for cell in table.rows[0].cells:
                cell_style = _get_cell_style(cell, ckeditor_mode)
                cell_content = _process_cell_text_with_images(cell, image_dict, printyn)
                
                # 헤더 셀에 워드 전용 속성 추가
                if cell_style:
                    # 기존 스타일에 워드 속성 추가
                    if 'style="' in cell_style:
                        cell_style = cell_style.replace('style="', 'style="mso-header-data: yes; ')
                    else:
                        cell_style += ' style="mso-header-data: yes;"'
                else:
                    cell_style = ' style="mso-header-data: yes;"'
                
                table_html.append(f'<th{cell_style}>{cell_content}</th>')
            
            table_html.append('</tr>')
            table_html.append('</thead>')
            
            # 바디 섹션 시작
            table_html.append('<tbody style="display: table-row-group;">')
            
            # 데이터 행 처리
            for row in table.rows[1:]:
                table_html.append('<tr>')
                for cell in row.cells:
                    cell_style = _get_cell_style(cell, ckeditor_mode)
                    cell_content = _process_cell_text_with_images(cell, image_dict, printyn)
                    table_html.append(f'<td{cell_style}>{cell_content}</td>')
                table_html.append('</tr>')
            
            table_html.append('</tbody>')
            
    except Exception as e:
        print(f"표 처리 중 오류 발생: {e}")
    
    table_html.append('</table>')
    return ''.join(table_html)


def _process_cell_text_with_images(cell, image_dict, printyn):
    """셀의 텍스트와 이미지를 HTML로 변환 (폰트 스타일 및 정렬 포함)"""
    result_parts = []
    
    try:
        for paragraph in cell.paragraphs:
            # 단락별로 처리하되, 각 Run의 스타일까지 적용
            para_parts = []
            current_width_px = 0
            
            # 단락의 정렬 스타일 확인
            alignment_style = _get_alignment_class(paragraph.alignment)
            indent_styles = _get_indent_styles(paragraph)
            
            # 탭 정지 위치 추출
            tab_stops = []
            try:
                if paragraph.paragraph_format.tab_stops:
                    for stop in paragraph.paragraph_format.tab_stops:
                        if stop.position:
                            tab_stops.append(_emu_to_px(stop.position))
                tab_stops.sort()
            except:
                pass
            
            for run in paragraph.runs:
                # 이미지 처리
                try:
                    blip_elements = run.element.xpath('.//a:blip')
                    for blip in blip_elements:
                        embed_attr = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_attr and embed_attr in image_dict:
                            image_data = image_dict[embed_attr]['data']
                            size_attr = _get_image_size_attr(run)
                            para_parts.append(f'<div class="image-container"><img src="{image_data}"{size_attr} alt="Document Image" /></div>')
                except:
                    pass
                
                # 텍스트 처리 (폰트 스타일 포함)
                processed_text = _process_run_text(run, tab_stops, current_width_px)
                if processed_text:
                    styled_text = _apply_text_styles(run, processed_text, printyn)
                    para_parts.append(styled_text)
            
            para_content = ''.join(para_parts)
            if para_content.strip():
                # 단락에 정렬 및 들여쓰기 스타일 적용
                all_styles = []
                if alignment_style:
                    # style 속성에서 스타일 값만 추출
                    style_content = alignment_style.split('"')[1] if '"' in alignment_style else alignment_style.replace('style=', '').strip()
                    all_styles.append(style_content)
                if indent_styles:
                    all_styles.extend(indent_styles)
                
                if all_styles:
                    style_attr = f' style="{";".join(all_styles)}"'
                    para_content = f'<div{style_attr}>{para_content}</div>'
                
                result_parts.append(para_content)
                
    except Exception as e:
        print(f"셀 처리 중 오류 발생: {e}")
        # 기본 처리 방식으로 폴백
        try:
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    result_parts.append(paragraph.text)
        except:
            pass
    
    return '<br>'.join(result_parts) if result_parts else ''


def _process_list(list_type, items):
    """리스트를 HTML로 변환"""
    if not items:
        return ''
    
    result = [f'<{list_type}>']
    for item in items:
        result.append(f'<li>{item}</li>')
    result.append(f'</{list_type}>')
    return ''.join(result)


def _emu_to_px(emu):
    """EMU(English Metric Unit)를 픽셀로 변환"""
    try:
        return int((emu / 914400) * 96)  # EMU → inches → px (96dpi 기준)
    except:
        return 0


def _get_cell_style(cell, use_ckeditor_format=False):
    """표 셀의 스타일(배경색, 테두리, 패딩 등) 추출"""
    style_attrs = []
    css_classes = []
    data_attrs = []
    
    try:
        # 셀 배경색 추출
        cell_elem = cell._element
        shd_elem = cell_elem.xpath('.//w:shd')
        
        if shd_elem:
            shd = shd_elem[0]
            fill_color = shd.get(qn('w:fill'))
            if fill_color and fill_color != 'auto' and len(fill_color) == 6:
                if use_ckeditor_format:
                    css_classes.append('cell-bg-color')
                    style_attrs.append(f'--cell-bg-color:#{fill_color}')
                    style_attrs.append(f'background-color:#{fill_color}')
                else:
                    style_attrs.append(f'background-color:#{fill_color}')
        
        # 셀 패딩 추가 (테이블 가독성 향상)
        style_attrs.append('padding:4px')
        style_attrs.append('border:1px solid black')
        
        # 셀의 주요 텍스트 정렬 확인 (모든 단락 확인)
        alignment_counts = {'left': 0, 'center': 0, 'right': 0, 'justify': 0}
        total_paragraphs = 0
        
        for para in cell.paragraphs:
            if para.text.strip():  # 빈 단락 제외
                total_paragraphs += 1
                alignment = para.alignment
                if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    alignment_counts['center'] += 1
                elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    alignment_counts['right'] += 1
                elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    alignment_counts['justify'] += 1
                else:
                    alignment_counts['left'] += 1
        
        # 가장 많이 사용된 정렬을 셀의 기본 정렬로 설정
        if total_paragraphs > 0:
            dominant_alignment = max(alignment_counts, key=alignment_counts.get)
            if dominant_alignment != 'left':  # left는 기본값이므로 생략
                if use_ckeditor_format:
                    if dominant_alignment == 'center':
                        css_classes.append('text-center')
                    elif dominant_alignment == 'right':
                        css_classes.append('text-right')
                    elif dominant_alignment == 'justify':
                        css_classes.append('text-justify')
                
                style_attrs.append(f'text-align:{dominant_alignment}')
        
        # 셀의 기본 폰트 정보 추출 (첫 번째 Run 기준)
        try:
            first_run = None
            for para in cell.paragraphs:
                if para.runs:
                    first_run = para.runs[0]
                    break
            
            if first_run:
                if first_run.font.size:
                    pt_size = first_run.font.size.pt * 1.33
                    style_attrs.append(f'font-size:{int(pt_size)}px')
                
                if first_run.font.name:
                    style_attrs.append(f"font-family:'{first_run.font.name}'")
                
                # 기본 굵기 확인
                if first_run.bold is True:
                    style_attrs.append('font-weight:bold')
                elif first_run.bold is None:
                    # 스타일에서 굵기 확인
                    try:
                        para_style = first_run._parent.style
                        while para_style:
                            if hasattr(para_style, "font") and para_style.font.bold:
                                style_attrs.append('font-weight:bold')
                                break
                            para_style = getattr(para_style, "based_on", None)
                    except:
                        pass
        except:
            pass
    
    except Exception as e:
        print(f"셀 스타일 처리 중 오류: {e}")
    
    # 결과 조합
    result_attrs = []
    if style_attrs:
        result_attrs.append(f'style="{";".join(style_attrs)}"')
    if css_classes:
        result_attrs.append(f'class="{" ".join(css_classes)}"')
    if data_attrs:
        result_attrs.extend(data_attrs)
    
    return f' {" ".join(result_attrs)}' if result_attrs else ''


def _rgb_to_hex(rgb_str):
    """RGB 문자열을 16진수로 변환"""
    try:
        if rgb_str and len(rgb_str) == 6:
            return f"#{rgb_str}"
    except:
        pass
    return None


# 기존 함수들과의 호환성을 위한 별칭
def emu_to_px(emu):
    return _emu_to_px(emu)

def get_alignment_class(alignment):
    return _get_alignment_class(alignment)

def process_paragraph_text(paragraph):
    """기본 텍스트 처리 (이미지 제외, 호환성 유지)"""
    return _process_paragraph_text_with_images(paragraph, {}, printyn=False)

def process_cell_text(cell):
    """기본 셀 텍스트 처리 (이미지 제외, 호환성 유지)"""
    return _process_cell_text_with_images(cell, {}, printyn=False)

def process_list(list_type, items):
    """기본 리스트 처리 (호환성 유지)"""
    return _process_list(list_type, items)