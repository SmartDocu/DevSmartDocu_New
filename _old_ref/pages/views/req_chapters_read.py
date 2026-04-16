import json, re, io, uuid, os, requests, time
from datetime import datetime, timezone, timedelta
from dateutil import parser
from django.conf import settings
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render, redirect
from django.views.decorators.csrf import csrf_exempt

from django.http import StreamingHttpResponse

from utilsPrj.supabase_client import get_supabase
# 워드 합치기
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 챕터 만들기
from utilsPrj.chapter_making import replace_doc
from utilsPrj.chapter_read import chapter_contents_read
from utilsPrj.html_to_docx import html_to_docx_merge    # 워드 만들기
from .req_doc_write import get_doc_name, get_chapter_name, add_page_number, add_total_pages, gendoc_save
from .req_doc_list import _build_params_value, _collect_required_columns, _filter_params_value_columns, _process_final_names, _create_finalnm_joined

#### 에러 로그 삽입 시 필요
from utilsPrj.errorlogs import error_log
import inspect

def req_chapters_read(request):
    # 세션 토큰
    supabase = get_supabase(request)

    user = request.session.get("user")
    if not user:
        return render(request, 'pages/home.html')
    
    if not user:
        # return JsonResponse({"result": "Failed", "message": "로그인이 필요합니다. 로그인 부탁드립니다."})
        # return redirect("login")
        code = 'login'
        text = '로그인이 필요합니다.'
        page = "req_read_chapters"
        return render(request, "pages/home.html", {
        "code": code,
        "text": text,
        "page": page,
        "request": request
    })
    user_id = user.get("id")
    
    ##### 0. docs 필터링
    # docs = supabase.schema("smartdoc").rpc("fn_docs_filtered__r_user", {'p_useruid': user_id}).execute().data or []
    # docid = [d["docid"] for d in docs]

    gendocuid = request.GET.get("gendocs")
    # print(f'GenDocUID: {gendocuid}')

    # 앞으로의 통일
    docid = user.get("docid")
    # print(docid)

    ##### 1. 작성 문서 목록
    # gendocs = supabase.schema("smartdoc").table("gendocs").select("*").in_("docid", docid).execute().data or []
    gendocs = supabase.schema("smartdoc").table("gendocs").select("*").eq("docid", docid).execute().data or []
    # print(f'GenDocs: {gendocs}')

    ##### 2. gendocs 의 정보 -> 매개변수 추출을 위함
    gendocs_data = supabase.schema('smartdoc').rpc("fn_gendocs__r", {"p_gendocuid": gendocuid}).execute().data or []
    # print(f'GenDocs_Data: {gendocs_data[0]}')

    ##### 3. gen 챕터
    genchapter_data = supabase.schema('smartdoc').rpc("fn_genchapters__r_gendocuid", {"p_gendocuid": gendocuid}).execute().data or []
    # print(f'GenChapter_Data: {genchapter_data}')

    doc_createfiledts = gendocs_data[0]['createfiledts']
    for i in gendocs_data:
        # 시간 포맷 정리
        if i.get('createfiledts'):
            try:
                dt = parser.parse(i['createfiledts']) if isinstance(i['createfiledts'], str) else i['createfiledts']
                i['createfiledts'] = dt.strftime("%y-%m-%d %H:%M")
            except Exception as e:
                i['createfiledts'] = ''

        if i.get('updatefiledts'):
            try:
                dt = parser.parse(i['updatefiledts']) if isinstance(i['updatefiledts'], str) else i['updatefiledts']
                i['updatefiledts'] = dt.strftime("%y-%m-%d %H:%M")
            except Exception as e:
                i['updatefiledts'] = ''

        if i.get('closedts'):
            try:
                dt = parser.parse(i['closedts']) if isinstance(i['closedts'], str) else i['closedts']
                i['closedts'] = dt.strftime("%y-%m-%d %H:%M")
            except Exception as e:
                i['closedts'] = ''

        if i.get('finaldts'):
            try:
                dt = parser.parse(i['finaldts']) if isinstance(i['finaldts'], str) else i['finaldts']
                finaldts = dt.strftime("%y-%m-%d %H:%M")
                if finaldts == '00-01-01 00:00':
                    i['finaldts'] = ''
                else:
                    i['finaldts'] = finaldts
            except Exception as e:
                i['finaldts'] = ''
        
        genparams = supabase.schema("smartdoc").rpc("fn_gendocs_params__r", {'p_gendocuid': i['gendocuid']}).execute().data or []
        i['params'] = json.dumps(genparams, ensure_ascii=False)

    # dataparams 조회
    dataparams = supabase.schema("smartdoc").table("dataparams").select("*").eq('docid', docid).order('orderno').execute().data or []
    
    data_ids = [d["datauid"] for d in dataparams if d.get("datauid") is not None]
    datas = supabase.schema("smartdoc").table("datas").select("*").in_("datauid", data_ids).execute().data

    # params_value 생성 (DataFrame을 dict 리스트로 변환)
    params_value = _build_params_value(request, docid, datas)

    # datauid별 필요한 컬럼 수집
    datauid_required_columns = _collect_required_columns(gendocs_data, dataparams)

    # params_value 컬럼 필터링
    _filter_params_value_columns(params_value, datauid_required_columns)

    # finalnm 생성
    _process_final_names(gendocs_data, params_value, dataparams)

    # finalnm_joined 생성
    _create_finalnm_joined(gendocs_data)

    for i in genchapter_data:
        # 시간 포맷 정리
        if i.get('createfiledts'):
            try:
                if doc_createfiledts:
                    i['new_chapteryn'] = i['createfiledts'] > doc_createfiledts
                else:
                    i['new_chapteryn'] = False

                dt = parser.parse(i['createfiledts']) if isinstance(i['createfiledts'], str) else i['createfiledts']
                i['createfiledts'] = dt.strftime("%y-%m-%d %H:%M")
            except Exception as e:
                print(e)
                i['createfiledts'] = ''
                i['new_chapteryn'] = False

        if i.get('updatefiledts'):
            try:
                if doc_createfiledts:
                    i['new_uploadyn'] = i['updatefiledts'] > doc_createfiledts
                else:
                    i['new_uploadyn'] = False

                dt = parser.parse(i['updatefiledts']) if isinstance(i['updatefiledts'], str) else i['updatefiledts']
                i['updatefiledts'] = dt.strftime("%y-%m-%d %H:%M")
            except Exception as e:
                i['updatefiledts'] = ''
                i['new_uploadyn'] = False

    # print(f'GenDocs: {gendocs_data}')
    # print(gendocs_data[0])
    # print(f'{genchapter_data}')
    
    return render(request, 'pages/req_chapters_read.html', {
        'docid': docid,
        'gendocuid': gendocuid,
        'gendocs': gendocs,
        'gendocs_data': gendocs_data[0],
        'genchapter_data': genchapter_data
        
    })

@csrf_exempt
def genchapter_file_upload(request):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'POST만 허용됨'})

    try:

        supabase = get_supabase(request)

        user = request.session.get("user")
        if not user:
            return JsonResponse({'success': False, 'message': '로그인 필요'})
        user_id = user.get("id")

        genchapteruid = request.POST.get("genchapteruid")
        gendocuid = request.POST.get("gendocuid")
        file_obj = request.FILES.get("file")

        if not file_obj:
            return JsonResponse({'success': False, 'message': '파일 없음'})

        # 1. 기존 파일 경로 확인
        # print('# 1. 기존 파일 경로 확인')
        old_data = supabase.schema("smartdoc").table("genchapters") \
            .select("updatefileurl, updatefilenm") \
            .eq("genchapteruid", genchapteruid).execute().data

        if old_data and old_data[0].get("updatefileurl"):
            # URL에서 파일 경로 추출 후 삭제
            # print('URL에서 파일 경로 추출 후 삭제')
            # old_path = f"{gendocuid}/chapters/{old_data[0]['updatefilenm']}"
            old_path = old_data[0]["updatefileurl"].split('smartdoc/')[1].rstrip('?')
            # print(old_path)
            supabase.storage.from_("smartdoc").remove([old_path])

        # 2. 새 파일 업로드
        # print('# 2. 새 파일 업로드')
        extension = os.path.splitext(file_obj.name)[1][1:]
        unique_filename = f"{uuid.uuid4()}.{extension}"
        path = f"result/{gendocuid}/chapters/{str(unique_filename)}"
        # InMemoryUploadedFile → bytes 변환
        file_bytes = file_obj.read()

        # 업로드 (upsert=True → 문자열로 변환)
        # print(f'Path: {path} / ')
        supabase.storage.from_("smartdoc").upload(
            path,
            file_bytes,
            {"cacheControl": "3600", "upsert": "true"}  # bool 대신 문자열
        )

        # 3. 퍼블릭 URL 가져오기
        # print('# 3. 퍼블릭 URL 가져오기')
        public_url = supabase.storage.from_("smartdoc").get_public_url(path)

        now = datetime.now().isoformat()
        # 4. DB 업데이트 1
        # print('# 4. DB 업데이트 1')
        supabase.schema("smartdoc").table("genchapters").update({
            "updatefileurl": public_url,
            "updatefilenm": file_obj.name,
            "updatefiledts": now,
            "updateuserid": user_id
        }).eq("genchapteruid", genchapteruid).execute()

        # 5. DB 업데이트 2
        # print('# 4. DB 업데이트 2')
        docid = supabase.schema("smartdoc").table("gendocs").select("*").eq("gendocuid", gendocuid).execute().data[0]['docid']
        supabase.schema("smartdoc").table("loguploads").insert({
            "objecttypenm": "C",
            "docid": docid,
            "gendocuid": gendocuid,
            "genchapteruid": genchapteruid,
            "updatefileurl": public_url,
            "updatefilenm": file_obj.name,
            "updatefiledts": now,
            "updateuserid": user_id
        }).execute()


        return JsonResponse({'success': True})

    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})


def chapter_detail_read(request):
    # print('chapter_detail_read 진입')
    # 세션 토큰
    supabase = get_supabase(request)

    user = request.session.get("user")
    
    body = json.loads(request.body)
    gendocuid = body.get("gendocuid")
    genchapteruid = body.get("genchapteruid")
    sep = body.get("sep")
    type = body.get("type")
    # print(f'Sep: {sep} / Type: {type} / GenChapterUID: {genchapteruid}')

    resp = chapter_contents_read(request, gendocuid, genchapteruid, sep, type)
    # print(resp['inmemoryyn'])
    contents = resp['contents']
    docyn = resp['docyn']
    autoyn = resp['autoyn']
    file_path = resp['file_path']
    file_name = resp['file_name']
    inmemoryyn = resp['inmemoryyn']

    if contents == '':
        contents = '작업된 항목이 없습니다.'

    if not gendocuid:
        gendocuid = supabase.schema('smartdoc').table('genchapters').select('*').eq('genchapteruid', genchapteruid).execute().data[0]['gendocuid']

    # 마감여부
    closeyn = supabase.schema("smartdoc").table("gendocs").select("*").eq("gendocuid", gendocuid).execute().data[0]['closeyn']
    # print({"contents": contents, "docyn": docyn, "autoyn": autoyn, "file_path": file_path, "file_name": file_name, "inmemoryyn": inmemoryyn})
    # print(f'File_Path: {file_path} / inmemory: {inmemoryyn}')

    return JsonResponse({"closeyn": closeyn, "contents": contents, "docyn": docyn, "autoyn": autoyn, "file_path": file_path, "file_name": file_name, "inmemoryyn": inmemoryyn, 'gendocuid': gendocuid, 'genchapteruid': genchapteruid})



@csrf_exempt
def chapter_rewrite(request):
    # print('Chapter_Rewrite 진입')
    supabase = get_supabase(request)
    user = request.session.get("user")
    user_id = user.get("id")

    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            genchapteruid = data.get("genchapteruid")
            
            if not genchapteruid:
                return JsonResponse({'success': False, 'message': '챕터 UID 없음'})

            now_dt = datetime.now() 
            now_iso = now_dt.isoformat()     # DB 저장용 (string)
            timeout = timedelta(hours=2)

            genchapters = supabase.schema('smartdoc').table('genchapters').select('*').eq('genchapteruid', genchapteruid).execute().data
            if not genchapters:
                return JsonResponse({'success': False, 'message': '해당 챕터를 찾을 수 없습니다.'})

            gendocuid = genchapters[0]["gendocuid"]

            # 1️ 문서 단위 락 조회 : 2시간 이상 진행 중인 row 강제 unlock / 내가 잡은 lock 강제 unlock
            genlocks_doc = (
                supabase
                .schema('smartdoc')
                .table('genlocks')
                .select('*')
                .eq('gendocuid', gendocuid)
                .eq('genchapteruid', '')  # 문서 락 row
                .eq('doclocked', True)
                .execute()
                .data
            )

            for lock in genlocks_doc:
                update_data = {}

                # 문서 락 체크
                start_time = datetime.fromisoformat(lock['docstartdts'])
                lockuser = lock['useruid']
                if now_dt - start_time > timeout:
                    update_data['doclocked'] = False
                    update_data['docenddts'] = now_iso
                if user_id == lockuser:
                    update_data['doclocked'] = False
                    update_data['docenddts'] = now_iso

                if update_data:
                    supabase.schema('smartdoc').table('genlocks').update(update_data)\
                        .eq('gendocuid', gendocuid)\
                        .eq('genchapteruid', '')\
                        .eq('doclocked', True)\
                        .execute()
                    
            # 2 챕터 단위 락 조회 : 2시간 이상 진행 중인 row 강제 unlock / 내가 잡은 lock 강제 unlock
            genlocks_chapter = (
                supabase
                .schema('smartdoc')
                .table('genlocks')
                .select('*')
                .eq('gendocuid', gendocuid)
                .eq('genchapteruid', genchapteruid)
                .eq('chapterlocked', True)
                .execute()
                .data
            )

            for lock in genlocks_chapter:
                update_data = {}

                # 문서 락 체크
                start_time = datetime.fromisoformat(lock['chapterstartdts'])
                lockuser = lock['useruid']
                if now_dt - start_time > timeout:
                    update_data['chapterlocked'] = False
                    update_data['chapterenddts'] = now_iso
                if user_id == lockuser:
                    update_data['chapterlocked'] = False
                    update_data['chapterenddts'] = now_iso

                if update_data:
                    supabase.schema('smartdoc').table('genlocks').update(update_data)\
                        .eq('gendocuid', gendocuid)\
                        .eq('genchapteruid', genchapteruid)\
                        .eq('chapterlocked', True)\
                        .execute()

            # 3️⃣ 남아있는 락 확인 (문서 락 OR 현재 챕터 락만)
            genlocks_c = (
                supabase
                .schema('smartdoc')
                .table('genlocks')
                .select('*')
                .eq('gendocuid', gendocuid)
                .eq('genchapteruid', genchapteruid)
                .eq('chapterlocked', True)
                .execute()
                .data
            )

            genlocks_d = (
                supabase
                .schema('smartdoc')
                .table('genlocks')
                .select('*')
                .eq('gendocuid', gendocuid)
                .eq('genchapteruid', '')  # 문서 락 row
                .eq('doclocked', True)
                .execute()
                .data
            )

            if genlocks_c + genlocks_d:
                def locked_stream():
                    yield f"data: {json.dumps({
                        'type': 'error',
                        'success': False,
                        'message': '이 문서 혹은 해당 챕터가 이미 작성 중입니다. 나중에 다시 시도해 주십시오.'
                    }, ensure_ascii=False)}\n\n"

                response = StreamingHttpResponse(
                    locked_stream(),
                    content_type='text/event-stream'
                )
                response['Cache-Control'] = 'no-cache'
                response['X-Accel-Buffering'] = 'no'
                return response

            try:
                # 4️⃣ 새 챕터 락 upsert
                supabase.schema('smartdoc').table('genlocks').upsert(
                    {
                        'gendocuid': gendocuid,
                        'genchapteruid': genchapteruid,  # 작업할 챕터
                        'doclocked': False,               # 문서 락 기록
                        'chapterlocked': True,           # 챕터 락 기록
                        'docstartdts': None,
                        'docenddts': None,
                        'chapterstartdts': now_iso,
                        'chapterenddts': None,
                        'useruid' : user_id,
                    },
                    on_conflict='gendocuid,genchapteruid'
                ).execute()
                
                def event_stream():
                    try:
                        for progress_data in replace_doc(request, supabase, user_id, genchapteruid, 'create', 'rewrite', 'Not',
                                                        genChapterDirectYn=True, divide="Chapter"):
                            yield f"data: {json.dumps(progress_data, ensure_ascii=False)}\n\n"
                        
                        yield f"data: {json.dumps({'type': 'complete', 'success': True}, ensure_ascii=False)}\n\n"
                        
                    except Exception as e:
                        error_msg = {'type': 'error', 'message': str(e)}
                        yield f"data: {json.dumps(error_msg, ensure_ascii=False)}\n\n"

            finally:
                # 스트림 종료 시 락 해제
                supabase.schema("smartdoc").table("genlocks").update({
                    "chapterlocked": False,
                    "chapterenddts": datetime.now().isoformat(),
                }).eq("gendocuid", gendocuid).eq("genchapteruid", genchapteruid).execute()


            response = StreamingHttpResponse(event_stream(), content_type='text/event-stream')
            response['Cache-Control'] = 'no-cache'
            response['X-Accel-Buffering'] = 'no'

            return response
        
        except Exception as e:
            print(f'ErrorMessage: {e}')
            return JsonResponse({'success': False, 'message': str(e)}, status=500)
    
    return JsonResponse({'success': False, 'message': 'POST 요청만 허용됨'})

@csrf_exempt
def doc_rewrite(request):
    # print('doc_rewrite 진입')

    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'POST 요청만 허용됨'})

    try:
        supabase = get_supabase(request)
        user = request.session.get("user")
        user_id = user.get("id")
        
        data = json.loads(request.body)
        gendocuid = data.get("gendocuid")
        if not gendocuid:
            return JsonResponse({'success': False, 'message': '문서 UID 없음'})

        doc = supabase.schema('smartdoc').table('gendocs').select("*").eq("gendocuid", gendocuid).execute().data[0]
        docid = doc['docid']
        gendocnm = doc['gendocnm']

        now_dt = datetime.now() 
        now_iso = now_dt.isoformat()     # DB 저장용 (string)
        timeout = timedelta(hours=2)

        # 1️⃣ 문서 기준 모든 락 조회
        genlocks = (
            supabase
            .schema('smartdoc')
            .table('genlocks')
            .select('*')
            .eq('gendocuid', gendocuid)
            .execute()
            .data
        )

        # 2️⃣ 2시간 이상 진행 중인 row 강제 unlock
        for lock in genlocks:
            update_data = {}

            # 문서 락
            if lock.get('doclocked') is True and lock.get('docstartdts'):
                start_time = datetime.fromisoformat(lock['docstartdts'])
                if user_id == lock['useruid'] or now_dt - start_time > timeout:
                    update_data['doclocked'] = False
                    update_data['docenddts'] = now_iso

            # 챕터 락
            if lock.get('chapterlocked') is True and lock.get('chapterstartdts'):
                start_time = datetime.fromisoformat(lock['chapterstartdts'])
                if user_id == lock['useruid'] or now_dt - start_time > timeout:
                    update_data['chapterlocked'] = False
                    update_data['chapterenddts'] = now_iso

            if update_data:
                supabase.schema('smartdoc').table('genlocks').update(update_data)\
                    .eq('gendocuid', lock['gendocuid'])\
                    .eq('genchapteruid', lock['genchapteruid'])\
                    .execute()

        # 3️⃣ 강제 unlock 후 남아있는 락 확인
        genlocks = (
            supabase
            .schema('smartdoc')
            .table('genlocks')
            .select('doclocked, chapterlocked')
            .eq('gendocuid', gendocuid)
            .execute()
            .data
        )

        is_locked = any(
            row.get('doclocked') is True or row.get('chapterlocked') is True
            for row in genlocks
        )

        if is_locked:
            # SSE로 락 메시지 전송하고 스트림 종료
            def locked_stream():
                yield f"data: {json.dumps({
                    'type': 'error',
                    'success': False,
                    'message': '문서가 작성 중에 있습니다. 나중에 다시 시도해 주십시오.'
                }, ensure_ascii=False)}\n\n"

            response = StreamingHttpResponse(
                locked_stream(),
                content_type='text/event-stream'
            )
            response['Cache-Control'] = 'no-cache'
            response['X-Accel-Buffering'] = 'no'
            return response


        # 4️⃣ 새 문서 락 upsert
        supabase.schema('smartdoc').table('genlocks').upsert(
            {
                'gendocuid': gendocuid,
                'genchapteruid': '',  # 문서 락 row
                'doclocked': True,
                'chapterlocked': False,
                'docstartdts': now_iso,
                'docenddts': None,
                'chapterstartdts': None,
                'chapterenddts': None,
            },
            on_conflict='gendocuid,genchapteruid'
        ).execute()

        chapters_response = supabase.schema("smartdoc").rpc("fn_genchapters__r_gendocuid", {'p_gendocuid': gendocuid}).execute()
        chapter_list = [ch["genchapteruid"] for ch in chapters_response.data]
        if not chapter_list:
            return JsonResponse({'success': False, 'message': '해당 문서에 챕터가 없습니다.'})

        loggendocuid = str(uuid.uuid4())
        update_loggendocs(supabase, loggendocuid, gendocuid, docid, user_id, 'str')

        def event_stream():
            try:
                total = len(chapter_list)
                last_sent = time.time()
                completed_chapters = []  # 완료된 챕터 추적

                # 1단계: 모든 챕터 처리
                for idx, genchapteruid in enumerate(chapter_list, 1):

                    if idx > 1:
                        yield f"data: {json.dumps({'type': 'wait', 'message': f'다음 챕터 준비 중...'}, ensure_ascii=False)}\n\n"
                        # time.sleep(10)

                    # print(f'[{idx}/{total}] GenChapterUID 처리 시작: {genchapteruid}')
                    try:
                        chapteruid = supabase.schema('smartdoc').table('genchapters').select('chapteruid').eq('genchapteruid', genchapteruid).execute().data[0]['chapteruid']
                        chapternm = supabase.schema('smartdoc').table('chapters').select('chapternm').eq('chapteruid', chapteruid).execute().data[0]['chapternm']
                        yield f"data: {json.dumps({'type': 'progress', 'obj': 'chapter', 'current': idx, 'total': total, 'message': f'챕터: {chapternm} 처리 시작 {idx}/{total}', 'status': 'processing'}, ensure_ascii=False)}\n\n"
                        
                        chapter_done = False
                        for progress_data in replace_doc(request, supabase, user_id, genchapteruid, 'create', 'rewrite', 'Not',
                                                        genChapterDirectYn=False, divide="Chapter", loggendocuid=loggendocuid, doc_write=True):
                            # progress_data에 챕터 정보 추가
                            if 'chapter_index' not in progress_data:
                                progress_data['chapter_index'] = idx
                                progress_data['chapter_total'] = total
                                progress_data['chapter_name'] = chapternm
                            
                            # print(f"Progress_Data: {progress_data['message']}")
                            return_message = progress_data['message']
                            chapter_message = f"챕터: {chapternm}\n"
                            progress_data['message'] = chapter_message + return_message
                            # texttemplate 초기화 -> 그냥 불러오면 글자 수 초과로 파싱 에러 발생
                            progress_data['texttemplate'] = ''
                            yield f"data: {json.dumps(progress_data, ensure_ascii=False)}\n\n"
                            
                            if progress_data.get('type') == 'progress':
                                if progress_data.get('explain') == '현재 챕터 생성 완료':
                                    # print(f'[{idx}/{total}] 챕터 완료 신호 수신: {genchapteruid} / {chapternm}')
                                    chapter_done = True
                                    completed_chapters.append(chapternm)
                                    # break 제거 - 제너레이터를 끝까지 소진

                            # heartbeat
                            if time.time() - last_sent > 3:
                                # yield ": heartbeat\n\n"
                                yield f"data: {json.dumps(progress_data, ensure_ascii=False)}\n\n"
                                last_sent = time.time()

                        if not chapter_done:
                            # print(f'[경고] 챕터 {idx} 완료 신호 누락: {genchapteruid}')
                            yield f"data: {json.dumps({'type': 'error', 'message': f'챕터 {chapternm} 완료 신호 누락'}, ensure_ascii=False)}\n\n"
                            continue
                        
                        # 챕터 완료 후 짧은 대기 (DB 저장 완료 보장)
                        time.sleep(1)
                        
                        # print(f'[{idx}/{total}] 챕터 처리 완료: {genchapteruid}')
                        # yield f"data: {json.dumps({'type': 'progress', 'obj': 'chapter', 'current': idx, 'total': total, 'message': f'챕터 {chapternm} 생성 완료'}, ensure_ascii=False)}\n\n"

                    except Exception as e:
                        yield f"data: {json.dumps({'type': 'error', 'obj': 'chapter', 'message': f'챕터 {chapternm} 처리 오류: {str(e)}'}, ensure_ascii=False)}\n\n"
                        continue

                # 모든 챕터 완료 확인
                # print(f'완료된 챕터 수: {len(completed_chapters)}/{total}')
                # print(f'완료된 챕터 UID: {completed_chapters}')
                
                if len(completed_chapters) != total:
                    yield f"data: {json.dumps({'type': 'error', 'obj': 'chapter', 'message': f'일부 챕터 처리 실패 ({len(completed_chapters)}/{total})'}, ensure_ascii=False)}\n\n"
                    return

                # 2단계: 문서 조합 시작 알림
                yield f"data: {json.dumps({'type': 'progress', 'obj': 'chapter_done', 'message': f'모든 챕터 완료. 문서 조합 시작...'}, ensure_ascii=False)}\n\n"
                time.sleep(1)  # 안정화 대기
                
                # 3단계: 문서 조합
                genchapter_list = [{'genchapteruid': ch} for ch in chapter_list]
                results = genchapter_list
                current_step = 0
                comp_doc = Document()
                previous_yn = False
                current_yn = False

                yield f"data: {json.dumps({'type': 'progress', 'obj': 'doc', 'step': 0, 'message': '문서 병합 준비 중...'}, ensure_ascii=False)}\n\n"
                # print('문서 조합 시작')
                
                for i, chapter in enumerate(results, 1):
                    # if i > 1:
                    #     yield f"data: {json.dumps({'type': 'wait', 'message': f'다음 챕터 준비 중... (10초 대기)'}, ensure_ascii=False)}\n\n"
                    #     time.sleep(10)

                    current_step += 1
                    chapter_name = get_chapter_name(supabase, chapter['genchapteruid'])
                    # print(f'ChapterName: {chapter_name}')
                    yield f"data: {json.dumps({'type': 'progress', 'obj': 'doc', 'step': current_step, 'total': len(results), 'message': f'챕터 [{chapter_name}] 병합 중...'}, ensure_ascii=False)}\n\n"
                    
                    # replace_doc 결과 확인
                    response = None
                    for result in replace_doc(request, supabase, user_id, chapter['genchapteruid'], 'create', 'write', 'Not',
                                            genChapterDirectYn=False, divide="Doc"):
                        if result.get('type') == 'complete':
                            response = result.get('texttemplate')
                            # break

                    if not response:
                        yield f"data: {json.dumps({'type': 'error', 'obj': 'doc', 'message': f'챕터 {chapter_name} 데이터 로드 실패'}, ensure_ascii=False)}\n\n"
                        continue

                    # HTML -> DOCX 병합
                    previous_yn, current_yn = html_to_docx_merge(supabase, comp_doc, chapter['genchapteruid'], response, current_step, previous_yn, current_yn)
                
                # 페이지 번호 추가
                yield f"data: {json.dumps({'type': 'progress', 'obj': 'doc', 'step': current_step + 1, 'message': '페이지 번호 추가 중...'}, ensure_ascii=False)}\n\n"
                for section in comp_doc.sections:
                    footer = section.footer
                    if not footer.paragraphs:
                        p = footer.add_paragraph("Page ")
                    else:
                        p = footer.paragraphs[0]
                        p.add_run(" | Page ")
                    add_page_number(p)
                    p.add_run(" / ")
                    add_total_pages(p)

                # 파일 업로드
                yield f"data: {json.dumps({'type': 'progress', 'obj': 'doc', 'step': current_step + 2, 'message': '파일 업로드 중...'}, ensure_ascii=False)}\n\n"
                filenm = f"{uuid.uuid4()}.docx"
                path = f"result/{gendocuid}/{filenm}"
                bucket_name = 'smartdoc'

                old_data = supabase.schema("smartdoc").table("gendocs") \
                                    .select("createfileurl") \
                                    .eq("gendocuid", gendocuid).execute().data

                if old_data and old_data[0].get("createfileurl"):
                    old_path = old_data[0]["createfileurl"].split('smartdoc/')[1].rstrip('?')
                    supabase.storage.from_("smartdoc").remove([old_path])
                
                buffer = io.BytesIO()
                comp_doc.save(buffer)
                buffer.seek(0)
                
                print(f'Bucket_name: {bucket_name} / Path: {path}')
                supabase.storage.from_(bucket_name).upload(path, buffer.read(), {"cacheControl": "3600", "upsert": "true"})
                
                public_url = supabase.storage.from_("smartdoc").get_public_url(path)

                # supabase = get_supabase(request)

                gendocs = {
                    'gendocuid': gendocuid,
                    'createfileurl': public_url,
                    'createfiledts': datetime.now().isoformat(),
                    'createuserid': user_id
                }
                gendoc_save(supabase, gendocs, gendocuid)

                gendocname = get_doc_name(supabase, gendocuid)
                
                if gendocname:
                    yield f"data: {json.dumps({'type': 'progress', 'obj': 'doc', 'message': '최종 저장 완료'}, ensure_ascii=False)}\n\n"

                update_loggendocs(supabase, loggendocuid, gendocuid, docid, user_id, 'end')

                # 최종 완료 신호
                yield f"data: {json.dumps({'type': 'complete', 'success': True, 'message': f'문서 [{gendocnm}] 작성 완료', 'url': public_url}, ensure_ascii=False)}\n\n"
            finally:
                supabase.schema("smartdoc").table("genlocks").update({
                    "doclocked": False,
                    "docenddts": now_iso,          # 종료 시간
                }).eq("gendocuid", gendocuid).eq("genchapteruid", "").execute()

        response = StreamingHttpResponse(event_stream(), content_type='text/event-stream')
        response['Cache-Control'] = 'no-cache'
        response['X-Accel-Buffering'] = 'no'
        return response

    except Exception as e:
        # --------------------------
        #  로그 저장
        # --------------------------
        try:
            # print(f'Error Message: {e}')
            # (request, errormessage, errorobject, creator, remark1 null 가능, remark2 null 가능, remark3 null 가능)
            error_log(request,
                      e, 
                      inspect.currentframe().f_code.co_name, 
                      user_id)

        except Exception as log_err:
            # pass 삭제 ➜ 최소한 서버에 출력 + 다시 raise
            # print("🔥 로그 저장 중 오류:", log_err)
            raise log_err

        return JsonResponse({'success': False, 'message': f"문서 작성 중 오류 발생: {str(e)}"}, status=500)

def update_loggendocs(supabase, loggendocuid, gendocuid, docid, user_id, str_end):
    # print(f'Update_LogGenDoc: {loggendocuid}')
    loggenchapter = None

    loggendoc = {
        'loggendocuid': loggendocuid,
    }

    if str_end == 'str':
        loggendoc['gendocuid'] = gendocuid
        loggendoc['docid'] = docid
        loggendoc['creator'] = user_id
        loggendoc['startdts'] = datetime.now().isoformat()
    elif str_end == 'end':
        loggendoc['enddts'] = datetime.now().isoformat()
    
    supabase.schema('smartdoc').table('loggendocs').upsert(loggendoc).execute()

def make_doc(request, supabase, gendocuid, results, total_steps, user_id):
    # print('MakeDoc 진입')
    now = datetime.now().isoformat()

    # 작업을 위한 함수 생성
    def generate_progress():
        try:
            current_step = 0
            comp_doc = Document()
            previous_yn = False
            current_yn = False

            yield f"data: {json.dumps({'type': 'progress', 'step': 0, 'message': '문서 조합 준비 중...'}, ensure_ascii=False)}\n\n"
            
            for i, chapter in enumerate(results, 1):
                current_step += 1
                # 챕터명 조회
                chapter_name = get_chapter_name(supabase, chapter['genchapteruid'])
                yield f"data: {json.dumps({'type': 'progress', 'step': current_step, 'message': f'챕터 {chapter_name} 병합 중...'}, ensure_ascii=False)}\n\n"
                
                # replace_doc 결과 확인
                response = None
                for result in replace_doc(request, supabase, user_id, chapter['genchapteruid'], 'create', 'write', 'Not',
                                          genChapterDirectYn=False, divide="Doc"):
                    if result.get('type') == 'complete':
                        response = result.get('texttemplate')
                        break

                if not response:
                    raise Exception(f'챕터 {chapter_name} 처리 실패')

                # HTML -> DOCX 병합
                previous_yn, current_yn = html_to_docx_merge(supabase, comp_doc, chapter['genchapteruid'], response, current_step, previous_yn, current_yn)
            
            # 2025-08-25
            # 페이지 번호 추가
            yield f"data: {json.dumps({'type': 'progress', 'step': current_step + 1, 'message': '페이지 번호 추가 중...'}, ensure_ascii=False)}\n\n"
            for section in comp_doc.sections:
                footer = section.footer
                if not footer.paragraphs:
                    p = footer.add_paragraph("Page ")
                else:
                    p = footer.paragraphs[0]
                    p.add_run(" | Page ")
                add_page_number(p)
                p.add_run(" / ")
                add_total_pages(p)

            # 파일 업로드
            yield f"data: {json.dumps({'type': 'progress', 'step': current_step + 2, 'message': '파일 업로드 중...'}, ensure_ascii=False)}\n\n"
            filenm = f"{uuid.uuid4()}.docx"
            path = f"result/{gendocuid}/{filenm}"
            bucket_name = 'smartdoc'

            old_data = supabase.schema("smartdoc").table("gendocs") \
                                .select("createfileurl") \
                                .eq("gendocuid", gendocuid).execute().data

            if old_data and old_data[0].get("createfileurl"):
                old_path = old_data[0]["createfileurl"].split('smartdoc/')[1].rstrip('?')
                supabase.storage.from_("smartdoc").remove([old_path])
            
            # 파일을 바이트로 읽어서 업로드
            # print('파일 바이트로 읽기')
            # 1. 메모리 버퍼 준비
            buffer = io.BytesIO()
            # 2. docx 저장
            comp_doc.save(buffer)
            # 3. 커서 처음으로 이동
            buffer.seek(0)

            # supabase storage upload
            # print('Supabase Upload')
            supabase.storage.from_(bucket_name).upload(path, buffer.read(), {"cacheControl": "3600", "upsert": "true"})
            
            # 공개 URL 확인
            # print('공개 URL 확인')
            public_url = supabase.storage.from_("smartdoc").get_public_url(path)

            gendocs = {
                'gendocuid': gendocuid
                ,'createfileurl': public_url
                ,'createfiledts': now
                ,'createuserid': user_id
            }
            gendoc_save(supabase, gendocs, gendocuid)

            # 문서명
            gendocname = get_doc_name(supabase, gendocuid)
            # 문서 저장
            # comp_doc.save(f"{gendocname}.docx")
            # 완료
            
            yield f"data: {json.dumps({'type': 'progress', 'message': '업로드 완료. 저장 중...'}, ensure_ascii=False)}\n\n"

            # yield f"data: {json.dumps({'type': 'complete', 'success': True, 'message': f'문서 {gendocname}의 조합이 완료되었습니다.'}, ensure_ascii=False)}\n\n"

        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': f'문서 조합 단계 오류: {str(e)}'}, ensure_ascii=False)}\n\n"


    return generate_progress()