## req_doc_write.py
import json, re, io, uuid, os, tempfile, subprocess
from io import BytesIO
from datetime import datetime, timezone, timedelta
from dateutil import parser
from django.conf import settings
from django.http import JsonResponse, FileResponse, HttpResponse, StreamingHttpResponse
from django.shortcuts import render, redirect
from django.views.decorators.csrf import csrf_exempt

# 워드 합치기
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from utilsPrj.supabase_client import get_supabase_client
from utilsPrj.chapter_making import replace_doc    # 챕터 만들기
from utilsPrj.html_to_docx import html_to_docx_merge    # 워드 만들기

#### 에러 로그 삽입 시 필요
from utilsPrj.errorlogs import error_log
import inspect


def req_doc_write(request):
    # 세션 토큰
    access_token = request.session.get("access_token")
    refresh_token = request.session.get("refresh_token")
    supabase = get_supabase_client(access_token, refresh_token)

    user = request.session.get("user")
    if not user:
        return render(request, 'pages/home.html')
    
    if not user:
        code = 'login'
        text = '로그인이 필요합니다.'
        page = "req_read_chapters"
        return render(request, "pages/home.html", {
        "code": code,
        "text": text,
        "page": page,
        "request": request
    })
    user_id = user.get("id")
    
    ##### 0. ㅈㄱ업을 위한 조회
    gendocuid = request.GET.get("gendocs")

    ##### 1. 작성 문서 목록
    gendocs = supabase.schema("smartdoc").table("gendocs").select("*").execute().data or []

    ##### 2. gendocs 의 정보 -> 매개변수 추출을 위함
    gendocs_data = supabase.schema('smartdoc').rpc("fn_gendocs__r", {"p_gendocuid": gendocuid}).execute().data or []

    ##### 3. gen 챕터
    genchapter_data = supabase.schema('smartdoc').rpc("fn_genchapters__r_gendocuid", {"p_gendocuid": gendocuid}).execute().data or []
    
    # 시간 포맷 정리
    for i in gendocs_data:
        if i.get('createfiledts'):
            try:
                dt = parser.parse(i['createfiledts']) if isinstance(i['createfiledts'], str) else i['createfiledts']
                i['createfiledts'] = dt.strftime("%y-%m-%d %H:%M")
            except Exception as e:
                i['createfiledts'] = ''

        if i.get('updatefiledts'):
            try:
                dt = parser.parse(i['updatefiledts']) if isinstance(i['updatefiledts'], str) else i['updatefiledts']
                i['updatefiledts'] = dt.strftime("%y-%m-%d %H:%M")
            except Exception as e:
                i['updatefiledts'] = ''
                
        if i.get('finaldts'):
            try:
                dt = parser.parse(i['finaldts']) if isinstance(i['finaldts'], str) else i['finaldts']
                finaldts = dt.strftime("%y-%m-%d %H:%M")
                if finaldts == '00-01-01 00:00':
                    i['finaldts'] = ''
                else:
                    i['finaldts'] = finaldts
            except Exception as e:
                i['finaldts'] = ''

    # 시간 포맷 정리
    for i in genchapter_data:
        if i.get('createfiledts'):
            try:
                dt = parser.parse(i['createfiledts']) if isinstance(i['createfiledts'], str) else i['createfiledts']
                i['createfiledts'] = dt.strftime("%y-%m-%d %H:%M")
            except Exception as e:
                i['createfiledts'] = ''

        if i.get('updatefiledts'):
            try:
                dt = parser.parse(i['updatefiledts']) if isinstance(i['updatefiledts'], str) else i['updatefiledts']
                i['updatefiledts'] = dt.strftime("%y-%m-%d %H:%M")
            except Exception as e:
                i['updatefiledts'] = ''
    
    return render(request, 'pages/req_doc_write.html', {
        'gendocuid': gendocuid,
        'gendocs': gendocs,
        'gendocs_data': gendocs_data[0],
        'genchapter_data': genchapter_data
        
    })

@csrf_exempt
def doc_write(request):
    # 세션 토큰
    access_token = request.session.get("access_token")
    refresh_token = request.session.get("refresh_token")
    supabase = get_supabase_client(access_token, refresh_token)

    user = request.session.get("user")
    user_id = user.get("id")

    # POST 신호 때만 동작
    if request.method == 'POST':
        try:
            # request 에서 필요 데이터 추출
            data = json.loads(request.body)
            results = data.get("results")
            gendocuid = data.get("gendocuid")

            now_dt = datetime.now() 
            now_iso = now_dt.isoformat()     # DB 저장용 (string)
            timeout = timedelta(hours=2)

            # 1️⃣ 문서 기준 모든 락 조회
            genlocks = (
                supabase
                .schema('smartdoc')
                .table('genlocks')
                .select('*')
                .eq('gendocuid', gendocuid)
                .execute()
                .data
            )

            # 2️⃣ 2시간 이상 진행 중인 row 강제 unlock
            for lock in genlocks:
                update_data = {}
                
                # 문서 락
                if lock.get('doclocked') is True and lock.get('docstartdts'):
                    start_time = datetime.fromisoformat(lock['docstartdts'])
                    if user_id == lock['useruid'] or now_dt - start_time > timeout:
                        update_data['doclocked'] = False
                        update_data['docenddts'] = now_iso

                # 챕터 락
                if lock.get('chapterlocked') is True and lock.get('chapterstartdts'):
                    start_time = datetime.fromisoformat(lock['chapterstartdts'])
                    if user_id == lock['useruid'] or now_dt - start_time > timeout:
                        update_data['chapterlocked'] = False
                        update_data['chapterenddts'] = now_iso


                if update_data:
                    supabase.schema('smartdoc').table('genlocks').update(update_data)\
                        .eq('gendocuid', lock['gendocuid'])\
                        .eq('genchapteruid', lock['genchapteruid'])\
                        .execute()

            # 3️⃣ 강제 unlock 후 남아있는 락 확인
            genlocks = (
                supabase
                .schema('smartdoc')
                .table('genlocks')
                .select('doclocked, chapterlocked')
                .eq('gendocuid', gendocuid)
                .execute()
                .data
            )

            is_locked = any(
                row.get('doclocked') is True or row.get('chapterlocked') is True
                for row in genlocks
            )

            if is_locked:
                # SSE 형식으로 "문서 작성 중" 메시지 바로 전달
                def locked_event():
                    yield f"data: {json.dumps({'type':'locked', 'message':'이 문서가 이미 작성 중입니다.'}, ensure_ascii=False)}\n\n"
                
                return StreamingHttpResponse(locked_event(), content_type='text/event-stream')

            # 4️⃣ 새 문서 락 upsert
            supabase.schema('smartdoc').table('genlocks').upsert(
                {
                    'gendocuid': gendocuid,
                    'genchapteruid': '',  # 문서 락 row
                    'doclocked': True,
                    'chapterlocked': False,
                    'docstartdts': now_iso,
                    'docenddts': None,
                    'chapterstartdts': None,
                    'chapterenddts': None,
                    'useruid' : user_id,
                },
                on_conflict='gendocuid,genchapteruid'
            ).execute()
            
            # 전체 단계 수 계산
            total_steps = len(results) + 3  # 챕터 수 + (HTML취합 + 작업 마무리 하는 중(Storage업로드 + 기존파일제거 + 업로드호출))

            # 작업을 위한 함수 생성
            def generate_progress():
                try:
                    current_step = 0
                    # 작업 시작 설명 처리
                    progress_data = {
                        'step': current_step, 
                        'total': total_steps, 
                        'message': f'작업 준비 중...', 
                        'status': 'processing'
                    }
                    yield f"data: {json.dumps(progress_data, ensure_ascii=False)}\n\n"
                    
                    try:
                        # 1. 챕터별 처리 단계
                        # Template = ''
                        comp_doc = Document()
                        # 이전 챕터 여부
                        previous_yn = False
                        # 현재 챕터 여부
                        current_yn = False
                        # 챕터 개수별 반복 진행
                        for i, chapter in enumerate(results, 1):
                            # print(f'GenChapterUID: {chapter["genchapteruid"]}')
                            current_step += 1
                            # 챕터명 조회
                            chapter_name = get_chapter_name(supabase, chapter["genchapteruid"])
                            
                            # 상태값 전달
                            progress_data = {
                                'step': current_step, 
                                'total': total_steps, 
                                'message': f'챕터: {chapter_name} 처리 중...', 
                                'status': 'processing'
                            }
                            yield f"data: {json.dumps(progress_data, ensure_ascii=False)}\n\n"
                            
                            # # HTML 코드 변환
                            # response = replace_doc(request, supabase, user_id, chapter["genchapteruid"], chapter["mode"], 'write', 'Not')
                            # Template += response
                            response = None
                            for result in replace_doc(request, supabase, user_id, chapter["genchapteruid"], chapter["mode"], 'write', 'Not',
                                                    genChapterDirectYn=False, divide="Doc"):
                                if result.get('type') == 'complete':
                                    response = result.get('texttemplate')
                                    break

                            if not response:
                                raise Exception(f'챕터 {i} 처리 실패') ####

                            # 챕터 이전 및 현재 Template 여부 확인
                            previous_yn, current_yn = html_to_docx_merge(supabase, comp_doc, chapter["genchapteruid"], response, current_step, previous_yn, current_yn)

                        # 2025-08-25
                        # 페이지 번호 추가
                        for section in comp_doc.sections:
                            footer = section.footer
                            # 기존 단락 없으면 새로 추가
                            if not footer.paragraphs:
                                p = footer.add_paragraph("Page ")
                            else:
                                p = footer.paragraphs[0]
                                p.add_run(" | Page ")

                            add_page_number(p)
                            p.add_run(" / ")
                            add_total_pages(p)

                        # 2. HTML 자료 취합
                        current_step += 1
                        progress_data = {
                            'step': current_step, 
                            'total': total_steps, 
                            'message': 'HTML 자료 취합 중...', 
                            'status': 'processing'
                        }
                        yield f"data: {json.dumps(progress_data, ensure_ascii=False)}\n\n"
                        
                        # 3 ~ 마지막은 마무리 단계 정리 중
                        current_step += 1
                        progress_data = {
                            'step': current_step, 
                            'total': total_steps, 
                            'message': '작업 정리 중', 
                            'status': 'processing'
                        }

                        # 3. 기존 파일 제거
                        # current_step += 1
                        # progress_data = {
                        #     'step': current_step, 
                        #     'total': total_steps, 
                        #     'message': '기존 파일 제거 중...', 
                        #     'status': 'processing'
                        # }
                        # yield f"data: {json.dumps(progress_data, ensure_ascii=False)}\n\n"
                        
                        bucket_name = 'smartdoc'
                        filenm = f"{uuid.uuid4()}.docx"
                        path = f"result/{gendocuid}/{filenm}"

                        old_data = supabase.schema("smartdoc").table("gendocs") \
                                            .select("createfileurl") \
                                            .eq("gendocuid", gendocuid).execute().data

                        if old_data and old_data[0].get("createfileurl"):
                            old_path = old_data[0]["createfileurl"].split('smartdoc/')[1].rstrip('?')
                            supabase.storage.from_("smartdoc").remove([old_path])
                        
                        # 파일을 바이트로 읽어서 업로드
                        # print('파일 바이트로 읽기')
                        # 1. 메모리 버퍼 준비
                        buffer = io.BytesIO()
                        # 2. docx 저장
                        comp_doc.save(buffer)
                        # 3. 커서 처음으로 이동
                        buffer.seek(0)

                        # supabase storage upload
                        # print('Supabase Upload')
                        supabase.storage.from_(bucket_name).upload(path, buffer.read(), {"cacheControl": "3600", "upsert": "true"})
                        
                        # 공개 URL 확인
                        # print('공개 URL 확인')
                        public_url = supabase.storage.from_("smartdoc").get_public_url(path)

                        gendocs = {
                            'gendocuid': gendocuid
                            ,'createfileurl': public_url
                            ,'createfiledts': datetime.now().isoformat()
                            ,'createuserid': user_id
                        }
                        gendoc_save(supabase, gendocs, gendocuid)

                        # 문서명
                        gendocname = get_doc_name(supabase, gendocuid)
                        # 문서 저장
                        # comp_doc.save(f"{gendocname}.docx")
                        # 완료
                        final_data = {
                            'step': total_steps, 
                            'total': total_steps, 
                            'message': '문서 작성 완료!', 
                            'status': 'completed', 
                            'path': path,
                            'docnm': gendocname
                        }
                        yield f"data: {json.dumps(final_data, ensure_ascii=False)}\n\n"
                        
                    except Exception as e:
                        print(f'ErrorMessage: {e}')
                        error_data = {
                            'step': 0, 
                            'total': total_steps, 
                            'message': f'', 
                            'status': 'error'
                        }
                        yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"

                finally:
                    end_time = datetime.now().isoformat()  # 종료 시점 반영
                    supabase.schema("smartdoc").table("genlocks").update({
                        "doclocked": False,
                        "docenddts": end_time,  # 실제 종료 시간
                    }).eq("gendocuid", gendocuid).eq("genchapteruid", "").execute()     

            return StreamingHttpResponse(generate_progress(), content_type='text/plain; charset=utf-8')

        except Exception as e:
            print(f'ErrorMessage: {e}')
            error_log(request, e, inspect.currentframe().f_code.co_name, user_id)

            return JsonResponse({'success': False, 'message': str(e)}, status=500)

    return JsonResponse({'success': False, 'message': 'POST 요청만 허용됨'})

def get_chapter_name(supabase, genchapteruid):
    """챕터 UID로 챕터명 조회"""
    try:
        genchapter = supabase.schema('smartdoc').table('genchapters').select('chapteruid').eq('genchapteruid', genchapteruid).execute().data
        chapteruid = genchapter[0]['chapteruid']
        chapter = supabase.schema('smartdoc').table('chapters').select('chapternm').eq('chapteruid', chapteruid).execute().data
        # print(chapter)
        return chapter[0]['chapternm'] if chapter else f"챕터 {genchapteruid[:8]}..."
    except:
        return f"챕터 {genchapteruid[:8]}..."

def get_doc_name(supabase, gendocuid):
    """문서 UID로 문서명 조회"""
    try:
        result = supabase.schema('smartdoc').table('gendocs').select('gendocnm').eq('gendocuid', gendocuid).execute().data
        return result[0]['gendocnm'] if result else f"챕터 {gendocuid[:8]}..."
    except:
        return f"챕터 {gendocuid[:8]}..."

def gendoc_save(supabase, gendocs, gendocuid):
    supabase.schema('smartdoc').table('gendocs').update(gendocs).eq('gendocuid', gendocuid).execute()

# 페이지 번호
def add_page_number(paragraph):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    run = OxmlElement('w:r')
    fldSimple.append(run)
    t = OxmlElement('w:t')
    t.text = " "
    run.append(t)
    paragraph._element.append(fldSimple)

# 페이지 전체체
def add_total_pages(paragraph):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'NUMPAGES')
    run = OxmlElement('w:r')
    fldSimple.append(run)
    t = OxmlElement('w:t')
    t.text = " "
    run.append(t)
    paragraph._element.append(fldSimple)