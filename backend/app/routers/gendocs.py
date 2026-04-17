import io
import json
import uuid
from datetime import datetime, timedelta
from typing import Optional
from urllib.parse import urlparse

from dateutil import parser as dp
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from backend.app.dependencies import get_token
from utilsPrj.supabase_client import get_thread_supabase, SUPABASE_SCHEMA

router = APIRouter()


def _sb(token: str):
    return get_thread_supabase(access_token=token)


def _get_user(token: str):
    sb = _sb(token)
    resp = sb.auth.get_user(token)
    if not resp or not resp.user:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="유효하지 않은 토큰입니다.")
    return resp.user


def _get_docid(sb, user_id: str) -> Optional[int]:
    row = sb.schema(SUPABASE_SCHEMA).table("users").select("mydocid").eq("useruid", user_id).execute().data
    docid = row[0].get("mydocid") if row else None
    return int(docid) if docid else None


def _fmt(val: Optional[str]) -> str:
    if not val:
        return ""
    try:
        dt = dp.parse(val) if isinstance(val, str) else val
        return dt.strftime("%y-%m-%d %H:%M")
    except Exception:
        return ""


class FakeRequest:
    """Minimal Django-like request stub for utilsPrj compatibility."""
    def __init__(self, access_token: str, user_id: str, docid: Optional[int] = None):
        self.session = {
            "access_token": access_token,
            "refresh_token": None,
            "user": {"id": user_id, "docid": str(docid) if docid else None},
        }
        self.method = "POST"


# ── Gendoc List ─────────────────────────────────────────────────────────────────

@router.get("")
def list_gendocs(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    token: str = Depends(get_token),
):
    user = _get_user(token)
    sb = _sb(token)
    user_id = str(user.id)
    docid = _get_docid(sb, user_id)
    if not docid:
        return {"gendocs": [], "docnm": None, "dataparams": []}

    today = datetime.now().date()
    sd = start_date or (today - timedelta(days=10)).strftime("%Y-%m-%d")
    ed = end_date or today.strftime("%Y-%m-%d")

    rpc_params = {"p_docid": docid, "p_start_date": sd}
    end_plus = (datetime.strptime(ed, "%Y-%m-%d") + timedelta(days=1)).strftime("%Y-%m-%d")
    rpc_params["p_end_date"] = end_plus

    rows = sb.schema(SUPABASE_SCHEMA).rpc("fn_gendocs__r_docid", rpc_params).execute().data or []
    docnm_resp = sb.schema(SUPABASE_SCHEMA).table("docs").select("docnm").eq("docid", docid).execute().data
    docnm = docnm_resp[0]["docnm"] if docnm_resp else None

    for item in rows:
        item["createfiledts"] = _fmt(item.get("createfiledts"))
        item["closedts"] = _fmt(item.get("closedts"))
        item["createdts"] = _fmt(item.get("createdts"))
        # params
        params = sb.schema(SUPABASE_SCHEMA).rpc("fn_gendocs_params__r", {"p_gendocuid": item["gendocuid"]}).execute().data or []
        item["params"] = params
        item["finalnm_joined"] = " / ".join(p.get("finalnm") or p.get("paramvalue") or "" for p in params if p.get("paramvalue"))

    dataparams = sb.schema(SUPABASE_SCHEMA).table("dataparams").select("*").eq("docid", docid).order("orderno").execute().data or []

    return {"gendocs": rows, "docnm": docnm, "dataparams": dataparams, "docid": docid}


# ── Gendoc Detail ───────────────────────────────────────────────────────────────

@router.get("/dataparams")
def get_dataparams(token: str = Depends(get_token)):
    user = _get_user(token)
    sb = _sb(token)
    docid = _get_docid(sb, str(user.id))
    if not docid:
        return {"dataparams": [], "params_value": []}

    dataparams = sb.schema(SUPABASE_SCHEMA).table("dataparams").select("*").eq("docid", docid).order("orderno").execute().data or []
    data_ids = [d["datauid"] for d in dataparams if d.get("datauid")]
    datas = sb.schema(SUPABASE_SCHEMA).table("datas").select("*").in_("datauid", data_ids).execute().data or [] if data_ids else []

    # Run each data source to get options
    from utilsPrj.process_data import process_data
    import pandas as pd, numpy as np
    from datetime import date

    def _convert(v):
        if isinstance(v, bytes):
            return v.decode("utf-8")
        if isinstance(v, (datetime, date)):
            return v.isoformat()
        if isinstance(v, (float,)) and (v != v):
            return None
        try:
            if isinstance(v, (np.integer, np.floating)):
                return v.item()
        except Exception:
            pass
        return v

    req = FakeRequest(token, str(user.id), docid)
    params_value = []
    for data_item in datas:
        try:
            df = process_data(req, datauid=data_item["datauid"], all=True)
            rows = df.to_dict("records") if not df.empty else []
            rows = [{k: _convert(v) for k, v in r.items()} for r in rows]
            rows = sorted(rows, key=lambda x: str(list(x.values())[0]) if x else "")
        except Exception:
            rows = []
        params_value.append({"datauid": data_item["datauid"], "value": rows})

    return {"dataparams": dataparams, "params_value": params_value}


@router.get("/{gendocuid}/status")
def get_gendoc_status(gendocuid: str, token: str = Depends(get_token)):
    _get_user(token)
    sb = _sb(token)
    status_rows = sb.schema(SUPABASE_SCHEMA).rpc("fn_gendoc_status__r", {"p_gendocuid": gendocuid}).execute().data or []
    gendoc = sb.schema(SUPABASE_SCHEMA).rpc("fn_gendocs__r", {"p_gendocuid": gendocuid}).execute().data or []

    for i in status_rows:
        i["createfiledts"] = _fmt(i.get("createfiledts"))
        i["updatefiledts"] = _fmt(i.get("updatefiledts"))

    gendocnm = gendoc[0]["gendocnm"] if gendoc else ""
    createfiledts = _fmt(gendoc[0].get("createfiledts")) if gendoc else ""
    return {"status": status_rows, "gendocnm": gendocnm, "createfiledts": createfiledts}


@router.get("/{gendocuid}/chapters")
def get_genchapters(gendocuid: str, token: str = Depends(get_token)):
    _get_user(token)
    sb = _sb(token)
    chapters = sb.schema(SUPABASE_SCHEMA).rpc("fn_genchapters__r_gendocuid", {"p_gendocuid": gendocuid}).execute().data or []
    for c in chapters:
        c["createfiledts"] = _fmt(c.get("createfiledts"))
        c["updatefiledts"] = _fmt(c.get("updatefiledts"))
    gendoc = sb.schema(SUPABASE_SCHEMA).rpc("fn_gendocs__r", {"p_gendocuid": gendocuid}).execute().data or []
    gendoc_info = gendoc[0] if gendoc else {}
    gendoc_info["createfiledts"] = _fmt(gendoc_info.get("createfiledts"))
    gendoc_info["updatefiledts"] = _fmt(gendoc_info.get("updatefiledts"))
    finaldts_raw = gendoc_info.get("finaldts")
    formatted_finaldts = _fmt(finaldts_raw)
    gendoc_info["finaldts"] = "" if formatted_finaldts == "00-01-01 00:00" else formatted_finaldts
    return {"chapters": chapters, "gendoc": gendoc_info}


# ── Gendoc Create ───────────────────────────────────────────────────────────────

class GendocCreateRequest(BaseModel):
    docid: int
    docnm: str
    params: list[dict]


@router.post("")
def create_gendoc(body: GendocCreateRequest, token: str = Depends(get_token)):
    user = _get_user(token)
    sb = _sb(token)
    user_id = str(user.id)
    now = datetime.now().isoformat()

    # 1. Create gendocs
    result = sb.schema(SUPABASE_SCHEMA).table("gendocs").insert({
        "docid": body.docid,
        "gendocnm": body.docnm,
        "creator": user_id,
    }).execute()
    gendocuid = result.data[0]["gendocuid"]

    # 2. Create gendoc_params
    if body.params:
        param_records = [
            {
                "gendocuid": gendocuid,
                "paramnm": p.get("paramnm"),
                "paramuid": p.get("paramuid"),
                "orderno": p.get("orderno"),
                "paramvalue": p.get("paramvalue"),
                "creator": user_id,
            }
            for p in body.params
        ]
        sb.schema(SUPABASE_SCHEMA).table("gendoc_params").insert(param_records).execute()

    # 3. Create genchapters (one per active chapter)
    chapters = (
        sb.schema(SUPABASE_SCHEMA).table("chapters")
        .select("*").eq("docid", body.docid).eq("useyn", True).execute().data or []
    )
    if chapters:
        genchapter_records = [
            {
                "docid": c["docid"],
                "chapteruid": c["chapteruid"],
                "gendocuid": gendocuid,
                "texttemplate": c.get("texttemplate"),
                "createfilestartdts": now,
                "creator": user_id,
                "createdts": now,
            }
            for c in chapters
        ]
        sb.schema(SUPABASE_SCHEMA).table("genchapters").insert(genchapter_records).execute()

    return {"gendocuid": gendocuid, "message": "생성되었습니다."}


# ── Gendoc Delete ───────────────────────────────────────────────────────────────

@router.delete("/{gendocuid}")
def delete_gendoc(gendocuid: str, token: str = Depends(get_token)):
    _get_user(token)
    sb = _sb(token)
    # Remove storage file
    row = sb.schema(SUPABASE_SCHEMA).table("gendocs").select("createfileurl").eq("gendocuid", gendocuid).execute().data
    if row and row[0].get("createfileurl"):
        url = row[0]["createfileurl"]
        parsed = urlparse(url)
        prefix = "/storage/v1/object/public/smartdoc/"
        if prefix in parsed.path:
            try:
                sb.storage.from_("smartdoc").remove([parsed.path.split(prefix)[-1]])
            except Exception:
                pass
    sb.schema(SUPABASE_SCHEMA).table("gendoc_params").delete().eq("gendocuid", gendocuid).execute()
    sb.schema(SUPABASE_SCHEMA).table("genchapters").delete().eq("gendocuid", gendocuid).execute()
    sb.schema(SUPABASE_SCHEMA).table("gendocs").delete().eq("gendocuid", gendocuid).execute()
    return {"message": "삭제되었습니다."}


# ── Params Update ────────────────────────────────────────────────────────────────

class GendocUpdateRequest(BaseModel):
    gendocuid: str
    gendocnm: str
    params: list[dict]


@router.post("/params/update")
def update_gendoc_params(body: GendocUpdateRequest, token: str = Depends(get_token)):
    _get_user(token)
    sb = _sb(token)
    sb.schema(SUPABASE_SCHEMA).table("gendocs").update({"gendocnm": body.gendocnm}).eq("gendocuid", body.gendocuid).execute()
    for p in body.params:
        sb.schema(SUPABASE_SCHEMA).table("gendoc_params").update({"paramvalue": p.get("paramvalue")}).eq("gendocuid", body.gendocuid).eq("paramuid", p.get("paramuid")).execute()
    return {"message": "파라미터가 변경되었습니다."}


# ── Params Check ─────────────────────────────────────────────────────────────────

class ParamsCheckRequest(BaseModel):
    docid: int
    params: list[dict]


@router.post("/params/check")
def check_params(body: ParamsCheckRequest, token: str = Depends(get_token)):
    _get_user(token)
    sb = _sb(token)
    if not body.params:
        return {"exists": False}

    paramuids = [p["paramuid"] for p in body.params if p.get("paramuid")]
    rows = sb.schema(SUPABASE_SCHEMA).table("gendoc_params").select("*").in_("paramuid", paramuids).execute().data or []

    from collections import defaultdict
    grouped = defaultdict(list)
    for r in rows:
        grouped[r["gendocuid"]].append(r)

    for gendocuid, items in grouped.items():
        if all(
            any(
                i["paramuid"] == p["paramuid"] and str(i["paramvalue"]).strip() == str(p["paramvalue"]).strip()
                for i in items
            )
            for p in body.params if p.get("paramuid")
        ):
            return {"exists": True, "gendocuid": gendocuid}

    return {"exists": False}


# ── Check Objects ────────────────────────────────────────────────────────────────

@router.post("/check-objects")
def check_objects(body: dict, token: str = Depends(get_token)):
    _get_user(token)
    sb = _sb(token)
    docid = body.get("docid")
    chapters = sb.schema(SUPABASE_SCHEMA).table("chapters").select("*").eq("docid", docid).execute().data or []
    unset = []
    for chap in chapters:
        objs = sb.schema(SUPABASE_SCHEMA).rpc("fn_objects__r", {"p_chapteruid": chap["chapteruid"]}).execute().data or []
        for obj in objs:
            if obj.get("useyn") and not obj.get("objectsettingyn"):
                unset.append({"text": f'챕터: {chap["chapternm"]} - 항목: {obj.get("objectnm", "")}'})
    return {"unset_objects": unset}


# ── Chapter Objects Read ──────────────────────────────────────────────────────────

@router.get("/genchapters/{genchapteruid}/objects")
def get_chapter_objects(genchapteruid: str, token: str = Depends(get_token)):
    _get_user(token)
    sb = _sb(token)

    genchap = sb.schema(SUPABASE_SCHEMA).table("genchapters").select("gendocuid,chapteruid,docid").eq("genchapteruid", genchapteruid).execute().data
    if not genchap:
        raise HTTPException(status_code=404, detail="챕터를 찾을 수 없습니다.")
    gendocuid = genchap[0]["gendocuid"]
    chapteruid = genchap[0]["chapteruid"]
    docid = genchap[0].get("docid")

    chapter = sb.schema(SUPABASE_SCHEMA).table("chapters").select("chapternm").eq("chapteruid", chapteruid).execute().data
    chapternm = chapter[0]["chapternm"] if chapter else ""

    gendoc = sb.schema(SUPABASE_SCHEMA).table("gendocs").select("gendocnm,closeyn").eq("gendocuid", gendocuid).execute().data
    closeyn = bool(gendoc[0]["closeyn"]) if gendoc else False
    gendocnm = gendoc[0]["gendocnm"] if gendoc else ""

    objects = sb.schema(SUPABASE_SCHEMA).rpc("fn_genobjects__r", {"p_genchapteruid": genchapteruid}).execute().data or []
    for obj in objects:
        obj["objcreatedts"] = _fmt(obj.get("objcreatedts"))
        obj["genobjcreatedts"] = _fmt(obj.get("genobjcreatedts"))
    objects = sorted(objects, key=lambda x: x.get("orderno", 0))

    return {
        "objects": objects,
        "chapternm": chapternm,
        "gendocnm": gendocnm,
        "gendocuid": gendocuid,
        "docid": docid,
        "chapteruid": chapteruid,
        "closeyn": closeyn,
    }


# ── Object Rewrite ───────────────────────────────────────────────────────────────

class ObjectRewriteRequest(BaseModel):
    objectuid: str


@router.post("/genchapters/{genchapteruid}/objects/{objectuid}/rewrite")
def rewrite_object(genchapteruid: str, objectuid: str, token: str = Depends(get_token)):
    user = _get_user(token)
    sb = _sb(token)
    user_id = str(user.id)
    docid = _get_docid(sb, user_id)

    genchap = sb.schema(SUPABASE_SCHEMA).table("genchapters").select("gendocuid").eq("genchapteruid", genchapteruid).execute().data
    if not genchap:
        raise HTTPException(status_code=404, detail="챕터를 찾을 수 없습니다.")
    gendocuid = genchap[0]["gendocuid"]

    # Check locks
    genlocks_c = sb.schema(SUPABASE_SCHEMA).table("genlocks").select("doclocked,chapterlocked").eq("gendocuid", gendocuid).eq("genchapteruid", genchapteruid).execute().data or []
    genlocks_d = sb.schema(SUPABASE_SCHEMA).table("genlocks").select("doclocked,chapterlocked").eq("gendocuid", gendocuid).eq("genchapteruid", "").execute().data or []
    is_locked = any(r.get("doclocked") or r.get("chapterlocked") for r in genlocks_c + genlocks_d)
    if is_locked:
        raise HTTPException(status_code=409, detail="이 문서의 해당 챕터가 이미 작성 중입니다.")

    req = FakeRequest(token, user_id, docid)

    try:
        from utilsPrj.chapter_making import replace_doc
        for progress_data in replace_doc(req, sb, user_id, genchapteruid, "create", "rewrite", objectuid, genObjectDirectYn=True):
            if progress_data.get("type") == "error":
                raise HTTPException(status_code=500, detail=progress_data.get("message", "오류가 발생했습니다."))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    return {"success": True}


# ── Apply Objects to Chapter ─────────────────────────────────────────────────────

@router.post("/genchapters/{genchapteruid}/apply")
def apply_chapter_objects(genchapteruid: str, token: str = Depends(get_token)):
    user = _get_user(token)
    sb = _sb(token)
    user_id = str(user.id)

    genchap = sb.schema(SUPABASE_SCHEMA).table("genchapters").select("gendocuid,chapteruid").eq("genchapteruid", genchapteruid).execute().data
    if not genchap:
        raise HTTPException(status_code=404, detail="챕터를 찾을 수 없습니다.")
    gendocuid = genchap[0]["gendocuid"]
    chapteruid = genchap[0]["chapteruid"]

    # Check locks
    genlocks_c = sb.schema(SUPABASE_SCHEMA).table("genlocks").select("doclocked,chapterlocked").eq("gendocuid", gendocuid).eq("genchapteruid", genchapteruid).execute().data or []
    genlocks_d = sb.schema(SUPABASE_SCHEMA).table("genlocks").select("doclocked,chapterlocked").eq("gendocuid", gendocuid).eq("genchapteruid", "").execute().data or []
    is_locked = any(r.get("doclocked") or r.get("chapterlocked") for r in genlocks_c + genlocks_d)
    if is_locked:
        raise HTTPException(status_code=409, detail="이 문서의 해당 챕터가 이미 작성 중입니다.")

    # Get base text template
    chapter_row = sb.schema(SUPABASE_SCHEMA).table("chapters").select("texttemplate").eq("chapteruid", chapteruid).execute().data
    texttemplate = chapter_row[0]["texttemplate"] if chapter_row else ""

    # Replace placeholders with generated object results
    read_texttemplate = sb.schema(SUPABASE_SCHEMA).rpc("fn_genchapter_detail__r", {"p_genchapteruid": genchapteruid}).execute().data or []
    for item in read_texttemplate:
        if item.get("genobjectuid"):
            read_datas = sb.schema(SUPABASE_SCHEMA).table("genobjects").select("resulttext").eq("genobjectuid", item["genobjectuid"]).execute().data
            html = read_datas[0]["resulttext"] if read_datas else ""
            placeholder = f"{{{{{item['objectnm']}}}}}"
            texttemplate = texttemplate.replace(placeholder, html or "")

    now = datetime.now().isoformat()

    # Update genchapters
    sb.schema(SUPABASE_SCHEMA).table("genchapters").update({
        "gentexttemplate": texttemplate,
        "genchapteruid": genchapteruid,
        "createuserid": user_id,
        "createfiledts": now,
    }).eq("genchapteruid", genchapteruid).execute()

    # Insert log
    try:
        sb.schema(SUPABASE_SCHEMA).table("gendoc_genchapters").insert({
            "gendocuid": gendocuid,
            "genchapteruid": genchapteruid,
            "creator": user_id,
            "createdts": now,
        }).execute()
    except Exception:
        pass

    return {"success": True, "message": "항목 반영이 완료되었습니다."}


# ── Full-Document Content Read (req_chapter_read 해당) ──────────────────────────

@router.get("/{gendocuid}/doc-content")
def get_doc_content(
    gendocuid: str,
    type: str = "auto",
    token: str = Depends(get_token),
):
    """전체 문서 HTML 내용 반환 (sep='doc') — Django chapter_read?sep=doc 에 해당"""
    user = _get_user(token)
    sb = _sb(token)
    user_id = str(user.id)
    docid = _get_docid(sb, user_id)
    req = FakeRequest(token, user_id, docid)

    try:
        from utilsPrj.chapter_read import chapter_contents_read
        resp = chapter_contents_read(req, gendocuid, None, "doc", type)
        contents = resp.get("contents") or "작업된 항목이 없습니다."
        file_path = resp.get("file_path")
        file_name = resp.get("file_name")
        inmemoryyn = resp.get("inmemoryyn", False)
    except Exception as e:
        contents = f"오류: {e}"
        file_path = None
        file_name = None
        inmemoryyn = False

    # 문서 정보 (작성자, 작성일시, 업로더, 업로드일시)
    doc_info = {}
    try:
        gd = sb.schema(SUPABASE_SCHEMA).table("gendocs").select("*").eq("gendocuid", gendocuid).execute().data
        if gd:
            d = gd[0]
            doc_info["gendocnm"] = d.get("gendocnm", "")
            doc_info["createfileurl"] = d.get("createfileurl")
            doc_info["updatefileurl"] = d.get("updatefileurl")
            doc_info["closeyn"] = bool(d.get("closeyn", False))
            # 작성자/업로더 이름
            for uid_field, nm_field, ts_field in [
                ("createuserid", "createuser", "createfiledts"),
                ("updateuserid", "updateuser", "updatefiledts"),
            ]:
                uid = d.get(uid_field)
                if uid:
                    try:
                        u = sb.schema("public").table("users").select("full_name").eq("useruid", uid).maybe_single().execute()
                        doc_info[nm_field] = u.data.get("full_name", "") if u.data else ""
                    except Exception:
                        doc_info[nm_field] = ""
                    ts = d.get(ts_field)
                    if ts:
                        try:
                            from dateutil import parser as dp
                            doc_info[ts_field] = dp.parse(ts).strftime("%y-%m-%d %H:%M")
                        except Exception:
                            doc_info[ts_field] = ts
    except Exception:
        pass

    return {
        "contents": contents,
        "file_path": file_path,
        "file_name": file_name,
        "inmemoryyn": inmemoryyn,
        "doc_info": doc_info,
        "type": type,
    }


# ── Chapter Content Read ─────────────────────────────────────────────────────────

@router.get("/genchapters/{genchapteruid}/content")
def get_chapter_content(
    genchapteruid: str,
    type: str = "auto",
    token: str = Depends(get_token),
):
    user = _get_user(token)
    sb = _sb(token)
    user_id = str(user.id)
    docid = _get_docid(sb, user_id)

    genchap = sb.schema(SUPABASE_SCHEMA).table("genchapters").select("gendocuid").eq("genchapteruid", genchapteruid).execute().data
    gendocuid = genchap[0]["gendocuid"] if genchap else None

    req = FakeRequest(token, user_id, docid)

    try:
        from utilsPrj.chapter_read import chapter_contents_read
        resp = chapter_contents_read(req, gendocuid, genchapteruid, "chapter", type)
        contents = resp.get("contents") or "작업된 항목이 없습니다."
        file_path = resp.get("file_path")
        file_name = resp.get("file_name")
        inmemoryyn = resp.get("inmemoryyn", False)
    except Exception as e:
        contents = f"오류: {e}"
        file_path = None
        file_name = None
        inmemoryyn = False

    closeyn = False
    if gendocuid:
        gd = sb.schema(SUPABASE_SCHEMA).table("gendocs").select("closeyn").eq("gendocuid", gendocuid).execute().data
        closeyn = bool(gd[0]["closeyn"]) if gd else False

    return {
        "contents": contents,
        "file_path": file_path,
        "file_name": file_name,
        "inmemoryyn": inmemoryyn,
        "closeyn": closeyn,
        "gendocuid": gendocuid,
        "genchapteruid": genchapteruid,
    }


# ── Chapter Rewrite (SSE) ────────────────────────────────────────────────────────

@router.post("/genchapters/{genchapteruid}/rewrite")
def rewrite_chapter(genchapteruid: str, token: str = Depends(get_token)):
    user = _get_user(token)
    sb = _sb(token)
    user_id = str(user.id)
    docid = _get_docid(sb, user_id)

    genchap = sb.schema(SUPABASE_SCHEMA).table("genchapters").select("gendocuid,chapteruid").eq("genchapteruid", genchapteruid).execute().data
    if not genchap:
        raise HTTPException(status_code=404, detail="챕터를 찾을 수 없습니다.")
    gendocuid = genchap[0]["gendocuid"]
    chapteruid = genchap[0]["chapteruid"]  # ← chapteruid도 함께 조회 필요

    def event_stream():
        from utilsPrj.chapter_making import replace_doc

        now_dt = datetime.now()
        now_iso = now_dt.isoformat()
        timeout = timedelta(hours=2)

        # Unlock stale locks
        genlocks = sb.schema(SUPABASE_SCHEMA).table("genlocks").select("*").eq("gendocuid", gendocuid).execute().data or []
        for lock in genlocks:
            upd = {}
            if lock.get("doclocked") and lock.get("docstartdts"):
                start = datetime.fromisoformat(lock["docstartdts"])
                if user_id == lock.get("useruid") or now_dt - start > timeout:
                    upd["doclocked"] = False
                    upd["docenddts"] = now_iso
            if lock.get("chapterlocked") and lock.get("chapterstartdts"):
                start = datetime.fromisoformat(lock["chapterstartdts"])
                if user_id == lock.get("useruid") or now_dt - start > timeout:
                    upd["chapterlocked"] = False
                    upd["chapterenddts"] = now_iso
            if upd:
                sb.schema(SUPABASE_SCHEMA).table("genlocks").update(upd).eq("gendocuid", gendocuid).eq("genchapteruid", lock["genchapteruid"]).execute()

        # Check remaining locks
        remaining = sb.schema(SUPABASE_SCHEMA).table("genlocks").select("doclocked,chapterlocked").eq("gendocuid", gendocuid).execute().data or []
        if any(r.get("doclocked") or r.get("chapterlocked") for r in remaining):
            yield f"data: {json.dumps({'type':'error','message':'이 문서 혹은 해당 챕터가 이미 작성 중입니다.'}, ensure_ascii=False)}\n\n"
            return

        # Set chapter lock
        sb.schema(SUPABASE_SCHEMA).table("genlocks").upsert({
            "gendocuid": gendocuid,
            "genchapteruid": genchapteruid,
            "doclocked": False,
            "chapterlocked": True,
            "docstartdts": None,
            "docenddts": None,
            "chapterstartdts": now_iso,
            "chapterenddts": None,
            "useruid": user_id,
        }, on_conflict="gendocuid,genchapteruid").execute()

        req = FakeRequest(token, user_id, docid)

        try:
            # ▼▼▼ 여기가 핵심 추가 부분 ▼▼▼
            from utilsPrj.template_parser import process_template, FunctionRegistry, extract_at_variables
            from utilsPrj.template_extracter import extract_from_processed_html

            # 1. 마스터 texttemplate 조회
            chapter_row = sb.schema(SUPABASE_SCHEMA).table("chapters").select("texttemplate").eq("chapteruid", chapteruid).execute().data
            texttemplate = chapter_row[0]["texttemplate"] if chapter_row else ""

            # 2. @변수 추출 및 context 빌드
            result = extract_at_variables(texttemplate)
            context = _build_context(sb, result["unique"])  # ← 아래 헬퍼 함수 필요

            # 3. 템플릿 처리 → flattexttemplate
            registry = FunctionRegistry()
            registry.set_default(lambda name, ctx, params: f"{{{{{name}}}}}[{json.dumps(params, ensure_ascii=False)}]")
            flattexttemplate = process_template(texttemplate, context, registry, True)

            # 4. genchapters에 flattexttemplate 저장
            sb.schema(SUPABASE_SCHEMA).table("genchapters").upsert({
                "genchapteruid": genchapteruid,
                "flattexttemplate": flattexttemplate,
            }).execute()

            # 5. genobjects 전체 갱신
            extracted = extract_from_processed_html(flattexttemplate)
            _upsert_genobjects(sb, extracted, genchapteruid, chapteruid, user_id)
            # ▲▲▲ 추가 끝 ▲▲▲

            # 기존 replace_doc 호출은 제거 또는 유지 (Django처럼 제거 권장)

            # for progress_data in replace_doc(req, sb, user_id, genchapteruid, "create", "rewrite", "Not",
            #                                  genChapterDirectYn=True, divide="Chapter"):
            #     yield f"data: {json.dumps(progress_data, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'type':'complete','success':True}, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type':'error','message':str(e)}, ensure_ascii=False)}\n\n"
        finally:
            sb.schema(SUPABASE_SCHEMA).table("genlocks").update({
                "chapterlocked": False,
                "chapterenddts": datetime.now().isoformat(),
            }).eq("gendocuid", gendocuid).eq("genchapteruid", genchapteruid).execute()

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ── Chapter File Upload ───────────────────────────────────────────────────────────

@router.post("/genchapters/{genchapteruid}/upload")
async def upload_chapter_file(
    genchapteruid: str,
    file: UploadFile = File(...),
    token: str = Depends(get_token),
):
    user = _get_user(token)
    sb = _sb(token)
    user_id = str(user.id)

    genchap = sb.schema(SUPABASE_SCHEMA).table("genchapters").select("gendocuid,updatefileurl").eq("genchapteruid", genchapteruid).execute().data
    if not genchap:
        raise HTTPException(status_code=404, detail="챕터를 찾을 수 없습니다.")
    gendocuid = genchap[0]["gendocuid"]

    if genchap[0].get("updatefileurl"):
        old_url = genchap[0]["updatefileurl"]
        parsed = urlparse(old_url)
        prefix = "/storage/v1/object/public/smartdoc/"
        if prefix in parsed.path:
            try:
                sb.storage.from_("smartdoc").remove([parsed.path.split(prefix)[-1]])
            except Exception:
                pass

    ext = file.filename.rsplit(".", 1)[-1] if "." in file.filename else "docx"
    path = f"result/{gendocuid}/chapters/{uuid.uuid4()}.{ext}"
    content = await file.read()
    sb.storage.from_("smartdoc").upload(path, content, {"cacheControl": "3600", "upsert": "true"})
    public_url = sb.storage.from_("smartdoc").get_public_url(path)
    now = datetime.now().isoformat()

    sb.schema(SUPABASE_SCHEMA).table("genchapters").update({
        "updatefileurl": public_url,
        "updatefilenm": file.filename,
        "updatefiledts": now,
        "updateuserid": user_id,
    }).eq("genchapteruid", genchapteruid).execute()

    docid_row = sb.schema(SUPABASE_SCHEMA).table("gendocs").select("docid").eq("gendocuid", gendocuid).execute().data
    if docid_row:
        sb.schema(SUPABASE_SCHEMA).table("loguploads").insert({
            "objecttypenm": "C",
            "docid": docid_row[0]["docid"],
            "gendocuid": gendocuid,
            "genchapteruid": genchapteruid,
            "updatefileurl": public_url,
            "updatefilenm": file.filename,
            "updatefiledts": now,
            "updateuserid": user_id,
        }).execute()

    return {"success": True, "message": "업로드되었습니다.", "url": public_url}


# ── File Upload ──────────────────────────────────────────────────────────────────

@router.post("/{gendocuid}/upload")
async def upload_file(
    gendocuid: str,
    file: UploadFile = File(...),
    token: str = Depends(get_token),
):
    user = _get_user(token)
    sb = _sb(token)
    user_id = str(user.id)

    old = sb.schema(SUPABASE_SCHEMA).table("gendocs").select("updatefileurl,updatefilenm").eq("gendocuid", gendocuid).execute().data
    if old and old[0].get("updatefileurl"):
        parsed = urlparse(old[0]["updatefileurl"])
        prefix = "/storage/v1/object/public/smartdoc/"
        if prefix in parsed.path:
            try:
                sb.storage.from_("smartdoc").remove([parsed.path.split(prefix)[-1]])
            except Exception:
                pass

    ext = file.filename.rsplit(".", 1)[-1] if "." in file.filename else "docx"
    path = f"result/{gendocuid}/{uuid.uuid4()}.{ext}"
    content = await file.read()
    sb.storage.from_("smartdoc").upload(path, content, {"cacheControl": "3600", "upsert": "true"})
    public_url = sb.storage.from_("smartdoc").get_public_url(path)
    now = datetime.now().isoformat()

    sb.schema(SUPABASE_SCHEMA).table("gendocs").update({
        "updatefileurl": public_url,
        "updatefilenm": file.filename,
        "updatefiledts": now,
        "updateuserid": user_id,
    }).eq("gendocuid", gendocuid).execute()

    docid = sb.schema(SUPABASE_SCHEMA).table("gendocs").select("docid").eq("gendocuid", gendocuid).execute().data
    if docid:
        sb.schema(SUPABASE_SCHEMA).table("loguploads").insert({
            "objecttypenm": "D",
            "docid": docid[0]["docid"],
            "gendocuid": gendocuid,
            "updatefileurl": public_url,
            "updatefilenm": file.filename,
            "updatefiledts": now,
            "updateuserid": user_id,
        }).execute()

    return {"message": "업로드되었습니다.", "url": public_url}


# ── Generate (SSE) ───────────────────────────────────────────────────────────────

class GenerateRequest(BaseModel):
    results: list[dict]


@router.post("/{gendocuid}/generate")
def generate_doc(gendocuid: str, body: GenerateRequest, token: str = Depends(get_token)):
    user = _get_user(token)
    sb = _sb(token)
    user_id = str(user.id)
    docid = _get_docid(sb, user_id)
    results = body.results

    def event_stream():
        from utilsPrj.chapter_making import replace_doc
        from utilsPrj.html_to_docx import html_to_docx_merge
        from docx import Document

        now_dt = datetime.now()
        now_iso = now_dt.isoformat()
        timeout = timedelta(hours=2)

        # Unlock stale locks
        genlocks = sb.schema(SUPABASE_SCHEMA).table("genlocks").select("*").eq("gendocuid", gendocuid).execute().data or []
        for lock in genlocks:
            upd = {}
            if lock.get("doclocked") and lock.get("docstartdts"):
                start = datetime.fromisoformat(lock["docstartdts"])
                if user_id == lock.get("useruid") or now_dt - start > timeout:
                    upd["doclocked"] = False
                    upd["docenddts"] = now_iso
            if lock.get("chapterlocked") and lock.get("chapterstartdts"):
                start = datetime.fromisoformat(lock["chapterstartdts"])
                if user_id == lock.get("useruid") or now_dt - start > timeout:
                    upd["chapterlocked"] = False
                    upd["chapterenddts"] = now_iso
            if upd:
                sb.schema(SUPABASE_SCHEMA).table("genlocks").update(upd).eq("gendocuid", gendocuid).eq("genchapteruid", lock["genchapteruid"]).execute()

        # Check remaining locks
        genlocks = sb.schema(SUPABASE_SCHEMA).table("genlocks").select("doclocked,chapterlocked").eq("gendocuid", gendocuid).execute().data or []
        if any(r.get("doclocked") or r.get("chapterlocked") for r in genlocks):
            yield f"data: {json.dumps({'type':'locked','message':'이 문서가 이미 작성 중입니다.'}, ensure_ascii=False)}\n\n"
            return

        # Set doc lock
        sb.schema(SUPABASE_SCHEMA).table("genlocks").upsert({
            "gendocuid": gendocuid,
            "genchapteruid": "",
            "doclocked": True,
            "chapterlocked": False,
            "docstartdts": now_iso,
            "docenddts": None,
            "chapterstartdts": None,
            "chapterenddts": None,
            "useruid": user_id,
        }, on_conflict="gendocuid,genchapteruid").execute()

        # 이전앱 동일: 챕터 수 + 3 (HTML취합 + 작업정리 + 완료)
        total_steps = len(results) + 3
        req = FakeRequest(token, user_id, docid)

        def _get_chap_name(genchapteruid, fallback):
            try:
                chapteruid = sb.schema(SUPABASE_SCHEMA).table("genchapters").select("chapteruid").eq("genchapteruid", genchapteruid).execute().data[0]["chapteruid"]
                resp = sb.schema(SUPABASE_SCHEMA).table("chapters").select("chapternm").eq("chapteruid", chapteruid).execute().data
                return resp[0]["chapternm"] if resp else fallback
            except Exception:
                return fallback

        def _add_page_number(paragraph):
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            fld = OxmlElement('w:fldSimple')
            fld.set(qn('w:instr'), 'PAGE')
            run = OxmlElement('w:r')
            fld.append(run)
            t = OxmlElement('w:t')
            t.text = " "
            run.append(t)
            paragraph._element.append(fld)

        def _add_total_pages(paragraph):
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            fld = OxmlElement('w:fldSimple')
            fld.set(qn('w:instr'), 'NUMPAGES')
            run = OxmlElement('w:r')
            fld.append(run)
            t = OxmlElement('w:t')
            t.text = " "
            run.append(t)
            paragraph._element.append(fld)

        try:
            current_step = 0
            yield f"data: {json.dumps({'step':current_step,'total':total_steps,'message':'작업 준비 중...','status':'processing'}, ensure_ascii=False)}\n\n"

            # ── 챕터별 처리: 이미 작성된 내용을 mode에 따라 읽어 DOCX 병합 ──────
            comp_doc = Document()
            previous_yn = False
            current_yn = False

            for i, chapter in enumerate(results, 1):
                current_step += 1
                genchapteruid = chapter["genchapteruid"]
                mode = chapter.get("mode", "create")  # 'create'=자동작성, 'update'=수정업로드
                chap_name = _get_chap_name(genchapteruid, f"챕터{i}")

                yield f"data: {json.dumps({'step':current_step,'total':total_steps,'message':f'챕터: {chap_name} 처리 중...','status':'processing'}, ensure_ascii=False)}\n\n"

                response = None
                for result in replace_doc(req, sb, user_id, genchapteruid, mode, "write", "Not",
                                          genChapterDirectYn=False, divide="Doc"):
                    if result.get("type") == "complete":
                        response = result.get("texttemplate")
                        break

                if not response:
                    raise Exception(f"챕터 {i} 처리 실패")

                previous_yn, current_yn = html_to_docx_merge(sb, comp_doc, genchapteruid, response, current_step, previous_yn, current_yn)

            # 페이지 번호 추가
            for section in comp_doc.sections:
                footer = section.footer
                if not footer.paragraphs:
                    p = footer.add_paragraph("Page ")
                else:
                    p = footer.paragraphs[0]
                    p.add_run(" | Page ")
                _add_page_number(p)
                p.add_run(" / ")
                _add_total_pages(p)

            # HTML 자료 취합 단계
            current_step += 1
            yield f"data: {json.dumps({'step':current_step,'total':total_steps,'message':'HTML 자료 취합 중...','status':'processing'}, ensure_ascii=False)}\n\n"

            # 작업 정리 단계 (Storage 업로드 + DB 갱신)
            current_step += 1
            yield f"data: {json.dumps({'step':current_step,'total':total_steps,'message':'작업 정리 중','status':'processing'}, ensure_ascii=False)}\n\n"

            filenm = f"{uuid.uuid4()}.docx"
            path = f"result/{gendocuid}/{filenm}"

            # Storage 작업은 service role 클라이언트 사용 (Django 이전앱과 동일하게 RLS 우회)
            from utilsPrj.supabase_client import get_service_client
            sb_svc = get_service_client()

            # 기존 파일 제거
            try:
                old = sb.schema(SUPABASE_SCHEMA).table("gendocs").select("createfileurl").eq("gendocuid", gendocuid).execute().data
                if old and old[0].get("createfileurl"):
                    old_url = old[0]["createfileurl"]
                    parsed = urlparse(old_url)
                    prefix = "/storage/v1/object/public/smartdoc/"
                    if prefix in parsed.path:
                        sb_svc.storage.from_("smartdoc").remove([parsed.path.split(prefix)[-1]])
            except Exception:
                pass  # 기존 파일 제거 실패는 무시

            # DOCX 메모리 저장
            buf = io.BytesIO()
            comp_doc.save(buf)
            buf.seek(0)

            # Storage 업로드
            try:
                sb_svc.storage.from_("smartdoc").upload(path, buf.read(), {"cacheControl": "3600", "upsert": "true"})
            except Exception as e:
                raise Exception(f"[Storage 업로드 오류] {e}")

            public_url = sb_svc.storage.from_("smartdoc").get_public_url(path)

            # gendocs DB 업데이트
            try:
                sb.schema(SUPABASE_SCHEMA).table("gendocs").update({
                    "createfileurl": public_url,
                    "createfiledts": datetime.now().isoformat(),
                    "createuserid": user_id,
                }).eq("gendocuid", gendocuid).execute()
            except Exception as e:
                raise Exception(f"[DB 업데이트 오류] {e}")

            gendocnm_resp = sb.schema(SUPABASE_SCHEMA).table("gendocs").select("gendocnm").eq("gendocuid", gendocuid).execute().data
            gendocnm = gendocnm_resp[0]["gendocnm"] if gendocnm_resp else ""

            yield f"data: {json.dumps({'step':total_steps,'total':total_steps,'message':'문서 작성 완료!','status':'completed','path':path,'docnm':gendocnm}, ensure_ascii=False)}\n\n"

        except Exception as e:
            yield f"data: {json.dumps({'step':0,'total':total_steps,'message':str(e),'status':'error'}, ensure_ascii=False)}\n\n"

        finally:
            sb.schema(SUPABASE_SCHEMA).table("genlocks").update({
                "doclocked": False,
                "docenddts": datetime.now().isoformat(),
            }).eq("gendocuid", gendocuid).eq("genchapteruid", "").execute()

    return StreamingResponse(event_stream(), media_type="text/event-stream")


def _build_context(sb, variables: list) -> dict:
    """req_chapters_read.py의 _build_context와 동일"""
    from utilsPrj.template_parser import parse_scalar_value
    context = {}
    for v in variables:
        find_tbl = sb.schema(SUPABASE_SCHEMA).table("tbl_param").select("*").eq("paramnm", v).execute().data
        if find_tbl:
            rows = []
            for row in find_tbl:
                raw = row.get("json")
                if not raw:
                    continue
                if isinstance(raw, str):
                    try:
                        parsed = json.loads(raw)
                        rows.extend(parsed) if isinstance(parsed, list) else rows.append(parsed)
                    except json.JSONDecodeError:
                        pass
                elif isinstance(raw, list):
                    rows.extend(raw)
                elif isinstance(raw, dict):
                    rows.append(raw)
            context[f"@{v}"] = rows
            continue

        find_sca = sb.schema(SUPABASE_SCHEMA).table("sca_param").select("*").eq("paramnm", v).execute().data
        if find_sca:
            context[f"@{v}"] = parse_scalar_value(find_sca[0]["value"])
    return context


def _upsert_genobjects(sb, extracted: list, genchapteruid: str, chapteruid: str, user_id: str):
    """req_chapters_read.py의 _upsert_genobjects와 동일"""
    now_iso = datetime.now().isoformat()
    sb.schema(SUPABASE_SCHEMA).table("genobjects").delete().eq("genchapteruid", genchapteruid).execute()

    rows = []
    for item in extracted:
        object_data = sb.schema(SUPABASE_SCHEMA).table("objects").select("*").eq("objectnm", item["objectNm"]).execute().data
        if not object_data:
            continue
        rows.append({
            "genobjectuid": str(uuid.uuid4()),
            "genchapteruid": genchapteruid,
            "chapteruid": chapteruid,
            "objectuid": object_data[0]["objectuid"],
            "objecttypecd": object_data[0].get("objecttypecd"),
            "filterjson": item["params"],
            "replacestring": item["replacestring"],
            "creator": user_id,
            "createdts": now_iso,
        })

    if rows:
        sb.schema(SUPABASE_SCHEMA).table("genobjects").insert(rows).execute()