import boto3
import io
import json
import logging
import threading
import uuid
from datetime import datetime, timezone
from typing import Optional
from urllib.parse import urlparse

from backend.app.config import settings
from utilsPrj.supabase_client import get_thread_supabase, get_service_client, SUPABASE_SCHEMA
from utilsPrj.chapter_making import replace_doc
from utilsPrj.credit_helper import apply_chapter_credit_deduction, apply_doc_credit_deduction
from utilsPrj.html_to_docx import html_to_docx_merge
from utilsPrj.notifications import create_notification


class FakeRequest:
    def __init__(self, access_token: str, user_id: str, docid: Optional[int] = None,
                 tenantid=None, projectid=None):
        self.session = {
            "access_token": access_token,
            "refresh_token": None,
            "user": {
                "id": user_id,
                "docid": str(docid) if docid else None,
                "tenantid": tenantid,
                "projectid": projectid,
            },
        }
        self.method = "POST"


def _build_context(sb, variables: list, req, docid) -> dict:
    from utilsPrj.process_data_ai import process_data_ai_preview
    context = {}
    for v in variables:
        find = sb.schema(SUPABASE_SCHEMA).table("datas").select("*").eq("datanm", v).eq("dfv_docid", docid).eq("datasourcecd", "dfv").execute().data
        if not find:
            continue
        sourcedatauid = find[0]["sourcedatauid"]
        datas = process_data_ai_preview(sb, req, sourcedatauid, find[0]["gensentence"], docid=docid)
        df = datas.get("result")
        if df is not None and not df.empty:
            datacols = sb.schema(SUPABASE_SCHEMA).table("datacols").select("querycolnm,dispcolnm").eq("datauid", sourcedatauid).execute().data
            disp_to_query = {item["dispcolnm"]: item["querycolnm"] for item in datacols}
            enriched = []
            for rec in df.to_dict("records"):
                row = dict(rec)
                for disp, query in disp_to_query.items():
                    if disp in rec:
                        row[query] = rec[disp]
                enriched.append(row)
            context[f"@{v}"] = enriched
        else:
            logger.warning("_build_context: @%s 데이터 비어있음 (sourcedatauid=%s)", v, sourcedatauid)
    return context


def _convert_filterjson_to_querycolnm(sb, dfv_datauids: list, raw_params: dict) -> dict:
    if not raw_params or not dfv_datauids:
        return raw_params
    ofm_resp = sb.schema(SUPABASE_SCHEMA).table("objectfiltermaps") \
        .select("dfvcolnm, objectdatauid, objectdatacolnm") \
        .in_("dfvdatauid", dfv_datauids).execute()
    if not ofm_resp.data:
        return raw_params
    dfvcolnm_to_ofm = {}
    for r in ofm_resp.data:
        if r["dfvcolnm"] not in dfvcolnm_to_ofm:
            dfvcolnm_to_ofm[r["dfvcolnm"]] = r
    obj_datauids = list({r["objectdatauid"] for r in ofm_resp.data if r.get("objectdatauid")})
    src_uid_map = {}
    if obj_datauids:
        datas_resp = sb.schema(SUPABASE_SCHEMA).table("datas") \
            .select("datauid, sourcedatauid").in_("datauid", obj_datauids).execute()
        src_uid_map = {r["datauid"]: r.get("sourcedatauid") for r in (datas_resp.data or [])}
    source_uids = [v for v in src_uid_map.values() if v]
    dc_map = {}
    if source_uids:
        dc_resp = sb.schema(SUPABASE_SCHEMA).table("datacols") \
            .select("datauid, querycolnm, dispcolnm").in_("datauid", source_uids).execute()
        for c in (dc_resp.data or []):
            dc_map.setdefault(c["datauid"], {})[c["dispcolnm"]] = c["querycolnm"]
    converted = {}
    for k, v in raw_params.items():
        ofm = dfvcolnm_to_ofm.get(k)
        if not ofm:
            converted[k] = v
            continue
        obj_disp = ofm["objectdatacolnm"]
        src_uid = src_uid_map.get(ofm.get("objectdatauid"))
        disp_col_map = dc_map.get(src_uid, {}) if src_uid else {}
        query_col = disp_col_map.get(obj_disp, obj_disp)
        converted[query_col] = v
    return converted


_DOMAIN_TBL_MAP = {"CU": "charts", "TU": "tables", "SU": "sentences",
                   "CA": "charts", "TA": "tables", "SA": "sentences"}


def _upsert_genobjects(sb, extracted: list, genchapteruid: str, chapteruid: str, user_id: str, docid=None,
                        projectid=None, tenantid=None, accountuid=None,
                        gendocjobuid=None, genchapterjobuid=None):
    now_iso = datetime.now(timezone.utc).isoformat()
    dfv_datauids = []
    if docid:
        dfv_rows = sb.schema(SUPABASE_SCHEMA).table("datas") \
            .select("datauid").eq("datasourcecd", "dfv").eq("dfv_docid", docid).execute().data
        dfv_datauids = [r["datauid"] for r in dfv_rows]

    # gencontenttypecd: 문서 전체 작성(D) > 챕터 작성(C) > 단일 항목 재작성(O)
    # 문서 전체 작성 시에는 gendocjobuid+genchapterjobuid 둘 다 들어오므로 gendocjobuid를 먼저 판별해야
    # 문서/챕터 단독 생성이 구분된다(gendocjobuid 없이 genchapterjobuid만 있으면 챕터 단독 생성).
    if gendocjobuid:
        gencontenttypecd = "D"
    elif genchapterjobuid:
        gencontenttypecd = "C"
    else:
        gencontenttypecd = "O"

    rows = []
    for item in extracted:
        obj = sb.schema(SUPABASE_SCHEMA).table("objects").select("*").eq("objectnm", item["objectNm"]).execute().data
        if not obj:
            continue
        objecttypecd = obj[0].get("objecttypecd")
        objectuid = obj[0]["objectuid"]
        filterjson = _convert_filterjson_to_querycolnm(sb, dfv_datauids, item["params"] or {})
        rows.append({
            "genobjectuid": str(uuid.uuid4()),
            "genchapteruid": genchapteruid,
            "chapteruid": chapteruid,
            "objectuid": obj[0]["objectuid"],
            "objecttypecd": obj[0].get("objecttypecd"),
            "filterjson": filterjson,
            "replacestring": item["replacestring"],
            "creator": user_id,
            "createdts": now_iso,
            "projectid": projectid,
            "tenantid": tenantid,
            "accountuid": accountuid,
            "gendocjobuid": gendocjobuid,
            "genchapterjobuid": genchapterjobuid,
            "gencontenttypecd": gencontenttypecd,
        })
    sb.schema(SUPABASE_SCHEMA).table("genobjects").delete().eq("genchapteruid", genchapteruid).execute()
    if rows:
        sb.schema(SUPABASE_SCHEMA).table("genobjects").insert(rows).execute()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
)
logger = logging.getLogger(__name__)

SQS_QUEUE_URL = settings.SQS_QUEUE_URL
SQS_CHAPTER_QUEUE_URL = settings.SQS_CHAPTER_QUEUE_URL
AWS_REGION = settings.AWS_REGION


def _update_queue(sb_svc, gendocjobuid, job_status_cd, *,
                  error_cd=None, error_msg=None, end_dts=None):
    row = {"jobstatuscd": job_status_cd}
    if error_cd is not None:
        row["errorcd"] = error_cd
    if error_msg is not None:
        row["errormessage"] = str(error_msg)[:2000]
    if end_dts is not None:
        row["enddts"] = end_dts
    sb_svc.schema(SUPABASE_SCHEMA).table("gendocs_realtimes").update(row).eq("gendocjobuid", gendocjobuid).execute()


def _update_chapter_queue(sb_svc, genchapterjobuid, job_status_cd, *,
                           error_cd=None, error_msg=None, end_dts=None):
    row = {"jobstatuscd": job_status_cd}
    if error_cd is not None:
        row["errorcd"] = error_cd
    if error_msg is not None:
        row["errormessage"] = str(error_msg)[:2000]
    if end_dts is not None:
        row["enddts"] = end_dts
    sb_svc.schema(SUPABASE_SCHEMA).table("genchapters_realtimes").update(row).eq("genchapterjobuid", genchapterjobuid).execute()


def _add_page_number(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    fld.append(run)
    t = OxmlElement("w:t")
    t.text = " "
    run.append(t)
    paragraph._element.append(fld)


def _add_total_pages(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "NUMPAGES")
    run = OxmlElement("w:r")
    fld.append(run)
    t = OxmlElement("w:t")
    t.text = " "
    run.append(t)
    paragraph._element.append(fld)


def _run_merge_and_upload(sb, sb_svc, req, gendocuid, docid, gendocnm, user_id, gendocjobuid, selected_chapters=None):
    """DOCX 병합 + Storage 업로드 + 완료 처리.

    selected_chapters(list[{genchapteruid, mode}])가 주어지면(문서 조합 작성) 그 챕터들만,
    각자 지정된 mode('create'=작성본 / 'update'=업로드본)로 병합한다.
    None이면(문서 전체 작성 fan-out 완료 후 호출) 기존 동작대로 gendocuid의 전체 챕터를
    모두 'create'(작성본) 기준으로 병합한다.
    """
    from docx import Document

    try:
        mode_by_uid = None
        if selected_chapters is not None:
            mode_by_uid = {c["genchapteruid"]: c.get("mode") or "create" for c in selected_chapters}

        # genchapters → chapters.chapterno 기준 정렬
        gc_rows = sb_svc.schema(SUPABASE_SCHEMA).table("genchapters") \
            .select("genchapteruid,chapteruid").eq("gendocuid", gendocuid).execute().data
        if mode_by_uid is not None:
            gc_rows = [r for r in gc_rows if r["genchapteruid"] in mode_by_uid]
        chapter_uids = [r["chapteruid"] for r in gc_rows]
        ch_rows = sb_svc.schema(SUPABASE_SCHEMA).table("chapters") \
            .select("chapteruid,chapterno").in_("chapteruid", chapter_uids).execute().data
        chapterno_map = {r["chapteruid"]: r["chapterno"] for r in ch_rows}
        sorted_gc = sorted(gc_rows, key=lambda r: chapterno_map.get(r["chapteruid"], 0))

        # Phase 2: DOCX 병합
        logger.info("Phase 2: DOCX 병합 (%s)", gendocuid)
        comp_doc = Document()
        previous_yn = current_yn = False

        for i, gc in enumerate(sorted_gc, 1):
            _genchapteruid = gc["genchapteruid"]
            make_type = mode_by_uid.get(_genchapteruid, "create") if mode_by_uid is not None else "create"
            response = None
            for result in replace_doc(req, sb, user_id, _genchapteruid, make_type, "write", "Not",
                                       genChapterDirectYn=False, divide="Doc"):
                if result.get("type") == "complete":
                    response = result.get("texttemplate")
                    break
            if not response:
                raise RuntimeError(f"챕터 {_genchapteruid} 데이터 로드 실패")
            previous_yn, current_yn = html_to_docx_merge(sb, comp_doc, _genchapteruid, response, i, previous_yn, current_yn)

        for section in comp_doc.sections:
            footer = section.footer
            if footer.paragraphs:
                p = footer.paragraphs[0]
                p.add_run(" | Page ")
            else:
                p = footer.add_paragraph("Page ")
            _add_page_number(p)
            p.add_run(" / ")
            _add_total_pages(p)

        # Phase 3: Storage 업로드
        logger.info("Phase 3: Storage 업로드 (%s)", gendocuid)
        filenm = f"{uuid.uuid4()}.docx"
        path = f"result/{gendocuid}/{filenm}"

        try:
            old = sb.schema(SUPABASE_SCHEMA).table("gendocs").select("createfileurl").eq("gendocuid", gendocuid).execute().data
            if old and old[0].get("createfileurl"):
                parsed = urlparse(old[0]["createfileurl"])
                prefix = "/storage/v1/object/public/sdoc/"
                if prefix in parsed.path:
                    sb_svc.storage.from_("sdoc").remove([parsed.path.split(prefix)[-1]])
        except Exception:
            pass

        buf = io.BytesIO()
        comp_doc.save(buf)
        buf.seek(0)
        sb_svc.storage.from_("sdoc").upload(path, buf.read(), {"cacheControl": "3600", "upsert": "true"})
        public_url = sb_svc.storage.from_("sdoc").get_public_url(path)

        sb.schema(SUPABASE_SCHEMA).table("gendocs").update({
            "createfileurl": public_url,
            "createfiledts": datetime.now(timezone.utc).isoformat(),
            "createuserid": user_id,
        }).eq("gendocuid", gendocuid).execute()

        end_iso = datetime.now(timezone.utc).isoformat()
        _update_queue(sb_svc, gendocjobuid, "E", end_dts=end_iso)
        logger.info("문서 완료: %s", gendocuid)

        try:
            apply_doc_credit_deduction(sb_svc, gendocjobuid)
        except Exception:
            logger.exception("크레딧 차감 실패: %s (gendocjobuid=%s)", gendocuid, gendocjobuid)

        create_notification(
            sb_svc, category="doc", status="info",
            title="문서 작성 완료", message=f"'{gendocnm}' 문서 작성이 완료되었습니다.",
            target_object="gendoc", target_uid=gendocuid, target_url="req/doc-read", target_useruid=user_id,
        )

        sb.schema(SUPABASE_SCHEMA).table("genlocks").update({
            "doclocked": False,
            "docenddts": datetime.now(timezone.utc).isoformat(),
        }).eq("gendocuid", gendocuid).eq("genchapteruid", "").execute()

    except Exception:
        logger.exception("Phase 2+3 오류: %s", gendocuid)
        try:
            import traceback
            end_iso = datetime.now(timezone.utc).isoformat()
            _update_queue(sb_svc, gendocjobuid, "E",
                          error_cd="ERR", error_msg=traceback.format_exc(), end_dts=end_iso)
            create_notification(
                sb_svc, category="doc", status="error",
                title="문서 작성 실패", message=f"'{gendocnm}' 문서 작성 중 오류가 발생했습니다.",
                target_object="gendoc", target_uid=gendocuid, target_url="req/doc-read", target_useruid=user_id,
            )
        except Exception:
            logger.exception("문서 오류 상태 업데이트 실패: %s", gendocuid)
        try:
            sb.schema(SUPABASE_SCHEMA).table("genlocks").update({
                "doclocked": False,
                "docenddts": datetime.now(timezone.utc).isoformat(),
            }).eq("gendocuid", gendocuid).eq("genchapteruid", "").execute()
        except Exception:
            logger.exception("잠금 해제 실패 (Phase 2+3 오류): %s", gendocuid)


def process_message(msg):
    body = json.loads(msg["Body"])
    gendocuid = body["gendocuid"]
    gendocjobuid = body["gendocjobuid"]
    user_id = body["user_id"]
    access_token = body["access_token"]
    docid = body.get("docid")
    tenantid = body.get("tenantid")
    projectid = body.get("projectid")
    accountuid = body.get("accountuid")
    gendocnm = body.get("gendocnm", "")
    receipt_handle = msg["ReceiptHandle"]

    sb = get_thread_supabase(access_token=access_token)
    sb_svc = get_service_client()
    sqs = boto3.client("sqs", region_name=AWS_REGION)

    # 문서 조합 작성 — 이미 작성된 챕터만 선택 그대로 병합 (Phase 1 fan-out/LLM 생성 없음 → 크레딧 차감 없음)
    if body.get("combine_only"):
        logger.info("문서 조합 작성 시작 (병합만, fan-out 없음): %s", gendocuid)
        req = FakeRequest(access_token, user_id, docid, tenantid=tenantid, projectid=projectid)
        try:
            _run_merge_and_upload(sb, sb_svc, req, gendocuid, docid, gendocnm, user_id, gendocjobuid,
                                   selected_chapters=body.get("chapters"))
        finally:
            try:
                sqs.delete_message(QueueUrl=SQS_QUEUE_URL, ReceiptHandle=receipt_handle)
            except Exception:
                logger.exception("SQS 메시지 삭제 실패 (조합 작성): %s", gendocuid)
        return

    results = body["results"]

    sb_svc.schema(SUPABASE_SCHEMA).table("gendocs_realtimes").update({
        "total_chapter_count": len(results),
    }).eq("gendocjobuid", gendocjobuid).execute()
    logger.info("처리 시작 (챕터 fan-out): %s, 총 %d개", gendocuid, len(results))

    start_iso = datetime.now(timezone.utc).isoformat()
    fan_out_success = False
    try:
        for chapter in results:
            genchapteruid = chapter["genchapteruid"]
            _gc = sb.schema(SUPABASE_SCHEMA).table("genchapters").select("chapteruid").eq("genchapteruid", genchapteruid).execute().data
            chapteruid = _gc[0]["chapteruid"] if _gc else ""
            res = sb_svc.schema(SUPABASE_SCHEMA).table("genchapters_realtimes").insert({
                "genchapteruid": genchapteruid,
                "docid": docid,
                "chapteruid": chapteruid,
                "jobstatuscd": "S",
                "startdts": start_iso,
                "creator": user_id,
                "is_start_doc": True,
                "gendocjobuid": gendocjobuid,
            }).execute()
            genchapterjobuid = res.data[0]["genchapterjobuid"]
            sb_svc.schema(SUPABASE_SCHEMA).table("genchapters").update({
                "genchapterjobuid": genchapterjobuid,
            }).eq("genchapteruid", genchapteruid).execute()
            sqs.send_message(
                QueueUrl=SQS_CHAPTER_QUEUE_URL,
                MessageBody=json.dumps({
                    "genchapteruid": genchapteruid,
                    "genchapterjobuid": genchapterjobuid,
                    "gendocuid": gendocuid,
                    "gendocjobuid": gendocjobuid,
                    "chapteruid": chapteruid,
                    "docid": docid,
                    "tenantid": tenantid,
                    "projectid": projectid,
                    "accountuid": accountuid,
                    "user_id": user_id,
                    "access_token": access_token,
                    "gendocnm": gendocnm,
                    "is_start_doc": True,
                }, ensure_ascii=False),
            )
            logger.info("챕터 큐 전송: %s", genchapteruid)

        fan_out_success = True
        logger.info("챕터 fan-out 완료: %s (%d개)", gendocuid, len(results))

    except Exception:
        logger.exception("fan-out 오류: %s", gendocuid)
        try:
            import traceback
            end_iso = datetime.now(timezone.utc).isoformat()
            _update_queue(sb_svc, gendocjobuid, "E",
                          error_cd="ERR", error_msg=traceback.format_exc(), end_dts=end_iso)
            create_notification(
                sb_svc, category="doc", status="error",
                title="문서 작성 실패", message=f"'{gendocnm}' 문서 작성 중 오류가 발생했습니다.",
                target_object="gendoc", target_uid=gendocuid, target_url="req/doc-read", target_useruid=user_id,
            )
        except Exception:
            logger.exception("큐 상태 업데이트 실패: %s", gendocuid)

    finally:
        if not fan_out_success:
            try:
                sb.schema(SUPABASE_SCHEMA).table("genlocks").update({
                    "doclocked": False,
                    "docenddts": datetime.now(timezone.utc).isoformat(),
                }).eq("gendocuid", gendocuid).eq("genchapteruid", "").execute()
            except Exception:
                logger.exception("잠금 해제 실패: %s", gendocuid)
        try:
            sqs.delete_message(QueueUrl=SQS_QUEUE_URL, ReceiptHandle=receipt_handle)
        except Exception:
            logger.exception("SQS 메시지 삭제 실패: %s", gendocuid)


def process_chapter_message(msg):
    body = json.loads(msg["Body"])
    genchapteruid = body["genchapteruid"]
    genchapterjobuid = body["genchapterjobuid"]
    gendocuid = body["gendocuid"]
    gendocjobuid = body.get("gendocjobuid")
    chapteruid = body["chapteruid"]
    docid = body.get("docid")
    tenantid = body.get("tenantid")
    projectid = body.get("projectid")
    accountuid = body.get("accountuid")
    user_id = body["user_id"]
    access_token = body["access_token"]
    gendocnm = body.get("gendocnm", "")
    is_start_doc = body.get("is_start_doc", False)
    receipt_handle = msg["ReceiptHandle"]

    sb = get_thread_supabase(access_token=access_token)
    sb_svc = get_service_client()
    sqs = boto3.client("sqs", region_name=AWS_REGION)

    chapternm = ""
    try:
        _chap_row = sb_svc.schema(SUPABASE_SCHEMA).table("chapters").select("chapternm").eq("chapteruid", chapteruid).execute().data
        chapternm = _chap_row[0]["chapternm"] if _chap_row else ""
    except Exception:
        pass

    logger.info("챕터 처리 시작: %s (is_start_doc=%s)", genchapteruid, is_start_doc)

    # SQS 재배달 중복 처리 방지 — 같은 genchapterjobuid는 최초 한 번만 'S'→'P' 선점 성공
    claim = sb_svc.schema(SUPABASE_SCHEMA).table("genchapters_realtimes") \
        .update({"jobstatuscd": "P"}) \
        .eq("genchapterjobuid", genchapterjobuid).eq("jobstatuscd", "S").execute()
    if not claim.data:
        logger.info("챕터 중복 처리 스킵 (이미 선점됨): %s", genchapteruid)
        try:
            sqs.delete_message(QueueUrl=SQS_CHAPTER_QUEUE_URL, ReceiptHandle=receipt_handle)
        except Exception:
            logger.exception("SQS 챕터 메시지 삭제 실패 (중복 스킵): %s", genchapteruid)
        return

    try:
        from utilsPrj.template_parser import process_template, FunctionRegistry, extract_at_variables
        from utilsPrj.template_extracter import extract_from_processed_html

        req = FakeRequest(access_token, user_id, docid, tenantid=tenantid, projectid=projectid)

        # 템플릿 처리 → flattexttemplate → genobjects
        _chap = sb.schema(SUPABASE_SCHEMA).table("chapters").select("texttemplate").eq("chapteruid", chapteruid).execute().data
        _tt = _chap[0]["texttemplate"] if _chap else ""
        _atv = extract_at_variables(_tt)
        _ctx = _build_context(sb, _atv["unique"], req, docid)
        _reg = FunctionRegistry()
        _reg.set_default(lambda name, ctx, params: f"{{{{{name}}}}}[{json.dumps(params, ensure_ascii=False)}]")
        _flat = process_template(_tt, _ctx, _reg, True)
        sb.schema(SUPABASE_SCHEMA).table("genchapters").upsert({"genchapteruid": genchapteruid, "flattexttemplate": _flat}).execute()
        _extracted = extract_from_processed_html(_flat)
        # 문서 전체 작성(fan-out) 시에는 gendocjobuid+genchapterjobuid 둘 다, 단일 챕터 작성 시에는
        # genchapterjobuid만 genobjects에 기록한다(gendocjobuid는 process_chapter_message 호출부에서
        # 이미 단독 작성 시 None으로 넘어옴 — 여기서 genchapterjobuid를 임의로 지우지 않는다).
        _upsert_genobjects(sb, _extracted, genchapteruid, chapteruid, user_id, docid=docid,
                           projectid=projectid, tenantid=tenantid, accountuid=accountuid,
                           gendocjobuid=gendocjobuid,
                           genchapterjobuid=genchapterjobuid)

        # LLM 콘텐츠 생성
        gen_chapter_direct = not is_start_doc
        for progress_data in replace_doc(req, sb, user_id, genchapteruid, "create", "rewrite", "Not",
                                          genChapterDirectYn=gen_chapter_direct, divide="Chapter",
                                          doc_write=is_start_doc,
                                          gendocjobuid=gendocjobuid,
                                          genchapterjobuid=genchapterjobuid):
            if progress_data.get("type") == "error":
                raise Exception(progress_data.get("message", "콘텐츠 생성 오류"))

        try:
            sb.schema(SUPABASE_SCHEMA).table("gendoc_genchapters").insert({
                "gendocuid": gendocuid,
                "genchapteruid": genchapteruid,
                "creator": user_id,
                "createdts": datetime.now(timezone.utc).isoformat(),
            }).execute()
        except Exception:
            pass

        end_iso = datetime.now(timezone.utc).isoformat()
        _update_chapter_queue(sb_svc, genchapterjobuid, "E", end_dts=end_iso)
        logger.info("챕터 완료: %s", genchapteruid)

        try:
            apply_chapter_credit_deduction(sb_svc, genchapterjobuid)
        except Exception:
            logger.exception("크레딧 차감 실패: %s (genchapterjobuid=%s)", genchapteruid, genchapterjobuid)

        if not is_start_doc:
            create_notification(
                sb_svc, category="chapter", status="info",
                title="챕터 작성 완료", message=f"'{gendocnm}' 문서의 '{chapternm}' 챕터 작성이 완료되었습니다.",
                target_object="gendoc", target_uid=gendocuid,
                target_url=f"req/chapters-read?genchapteruid={genchapteruid}", target_useruid=user_id,
            )

        # 문서 작성 fan-out 챕터: 전체 완료 여부 체크 → 마지막이면 Phase 2+3
        if is_start_doc and gendocjobuid:
            try:
                done_count = len(sb_svc.schema(SUPABASE_SCHEMA).table("genchapters_realtimes") \
                    .select("genchapterjobuid") \
                    .eq("gendocjobuid", gendocjobuid).eq("jobstatuscd", "E").execute().data)

                doc_rt = sb_svc.schema(SUPABASE_SCHEMA).table("gendocs_realtimes") \
                    .select("total_chapter_count,gendocnm,jobstatuscd").eq("gendocjobuid", gendocjobuid).execute().data
                total_cnt = doc_rt[0].get("total_chapter_count") if doc_rt else 0
                stored_gendocnm = doc_rt[0].get("gendocnm", gendocnm) if doc_rt else gendocnm

                logger.info("챕터 완료 체크: %s (%d/%d)", gendocuid, done_count, total_cnt)

                if total_cnt and done_count >= total_cnt:
                    claim = sb_svc.schema(SUPABASE_SCHEMA).table("gendocs_realtimes") \
                        .update({"jobstatuscd": "merging"}) \
                        .eq("gendocjobuid", gendocjobuid).eq("jobstatuscd", "S").execute()
                    if claim.data:
                        logger.info("Phase 2+3 시작: %s", gendocuid)
                        _run_merge_and_upload(sb, sb_svc, req, gendocuid, docid, stored_gendocnm, user_id, gendocjobuid)
                    else:
                        logger.info("Phase 2+3 이미 다른 워커가 처리 중: %s", gendocuid)
            except Exception:
                logger.exception("완료 체크 실패: %s", gendocuid)

    except Exception:
        logger.exception("챕터 오류: %s", genchapteruid)
        try:
            import traceback
            end_iso = datetime.now(timezone.utc).isoformat()
            _update_chapter_queue(sb_svc, genchapterjobuid, "E",
                                   error_cd="ERR", error_msg=traceback.format_exc(), end_dts=end_iso)
            if not is_start_doc:
                create_notification(
                    sb_svc, category="chapter", status="error",
                    title="챕터 작성 실패", message=f"'{gendocnm}' 문서의 '{chapternm}' 챕터 작성 중 오류가 발생했습니다.",
                    target_object="gendoc", target_uid=gendocuid,
                    target_url=f"req/chapters-read?genchapteruid={genchapteruid}", target_useruid=user_id,
                )
        except Exception:
            logger.exception("챕터 큐 상태 업데이트 실패: %s", genchapteruid)

        if is_start_doc and gendocjobuid:
            try:
                doc_rt = sb_svc.schema(SUPABASE_SCHEMA).table("gendocs_realtimes") \
                    .select("gendocnm,jobstatuscd").eq("gendocjobuid", gendocjobuid).execute().data
                if doc_rt and doc_rt[0].get("jobstatuscd") not in ("E", "merging"):
                    end_iso = datetime.now(timezone.utc).isoformat()
                    _update_queue(sb_svc, gendocjobuid, "E",
                                  error_cd="ERR", error_msg=f"챕터 처리 실패: {genchapteruid}", end_dts=end_iso)
                    create_notification(
                        sb_svc, category="doc", status="error",
                        title="문서 작성 실패",
                        message=f"'{doc_rt[0].get('gendocnm') or gendocnm}' 문서 작성 중 오류가 발생했습니다.",
                        target_object="gendoc", target_uid=gendocuid, target_url="req/doc-read", target_useruid=user_id,
                    )
                    sb.schema(SUPABASE_SCHEMA).table("genlocks").update({
                        "doclocked": False,
                        "docenddts": datetime.now(timezone.utc).isoformat(),
                    }).eq("gendocuid", gendocuid).eq("genchapteruid", "").execute()
            except Exception:
                logger.exception("문서 오류 처리 실패: %s", gendocuid)

    finally:
        try:
            sb.schema(SUPABASE_SCHEMA).table("genlocks").update({
                "chapterlocked": False,
                "chapterenddts": datetime.now(timezone.utc).isoformat(),
            }).eq("genchapteruid", genchapteruid).execute()
        except Exception:
            logger.exception("챕터 잠금 해제 실패: %s", genchapteruid)
        try:
            sqs.delete_message(QueueUrl=SQS_CHAPTER_QUEUE_URL, ReceiptHandle=receipt_handle)
        except Exception:
            logger.exception("SQS 챕터 메시지 삭제 실패: %s", genchapteruid)


def poll_queue(queue_url, handler_fn):
    sqs = boto3.client("sqs", region_name=AWS_REGION)
    logger.info("큐 폴링 시작 — %s", queue_url)
    while True:
        resp = sqs.receive_message(
            QueueUrl=queue_url,
            MaxNumberOfMessages=1,
            WaitTimeSeconds=20,
        )
        for msg in resp.get("Messages", []):
            try:
                handler_fn(msg)
            except Exception:
                logger.exception("메시지 처리 중 예외 [%s]", queue_url)


def main():
    # 큐 URL이 설정된 핸들러만 스레드 시작 (미설정 큐는 건너뜀)
    queue_handlers = {
        SQS_QUEUE_URL:         process_message,
        SQS_CHAPTER_QUEUE_URL: process_chapter_message,
    }
    threads = [
        threading.Thread(target=poll_queue, args=(url, fn), daemon=True, name=f"worker-{url.split('/')[-1]}")
        for url, fn in queue_handlers.items()
        if url
    ]
    if not threads:
        logger.error("활성 SQS 큐 URL이 없습니다. 환경변수를 확인하세요.")
        return
    for t in threads:
        t.start()
    logger.info("SmartDocu Worker 시작 — %d개 큐 폴링 중", len(threads))
    for t in threads:
        t.join()


if __name__ == "__main__":
    main()
